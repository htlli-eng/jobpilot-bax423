from __future__ import annotations

from collections import Counter
from pathlib import Path
import ast
import io
import json
import re

import numpy as np
import pandas as pd
import streamlit as st

from embedding_utils import (
    build_user_profile_text,
    embedding_files_exist,
    retrieve_top_candidates,
)
from resume_utils import (
    extract_candidate_name,
    extract_skills,
    extract_text_from_pdf,
    format_resume_preview,
)


PROJECT_ROOT = Path(__file__).resolve().parent
DATA_DIR = PROJECT_ROOT / "data"
SAMPLE_DATA_PATH = DATA_DIR / "jobs_sample.csv"
CLEANED_DATA_PATH = DATA_DIR / "jobs_cleaned.csv"
EMBEDDING_CANDIDATE_COUNT = 500
EMBEDDING_BONUS_SCALE = 30
SALARY_SLIDER_MAX = 250_000
SENIORITY_PREFERENCE_OPTIONS = [
    "No preference",
    "Entry-level / New Grad only",
    "Junior to Mid-level",
    "Mid to Senior allowed",
    "Senior+ preferred",
    "Exclude Senior/Staff/Principal",
    "Exclude Junior/Entry-level",
]
COMPANY_SIZE_PREFERENCE_OPTIONS = [
    "No preference",
    "Prefer small companies",
    "Prefer medium companies",
    "Prefer large companies",
    "Exclude tiny startups",
    "Large companies only",
]
TINY_COMPANY_SIGNALS = {
    "startup",
    "early-stage",
    "early stage",
    "seed-stage",
    "seed stage",
    "series a",
    "small team",
    "founding team",
    "under 50 employees",
    "under 100 employees",
    "fewer than 100 employees",
    "less than 100 employees",
}
LARGE_COMPANY_SIGNALS = {
    "fortune 500",
    "global company",
    "multinational",
    "enterprise",
    "public company",
    "large-scale",
    "large scale",
    "thousands of employees",
    "worldwide",
    "established company",
}
BAY_AREA_TERMS = {
    "bay area",
    "san francisco",
    "san jose",
    "oakland",
    "palo alto",
    "mountain view",
    "sunnyvale",
    "santa clara",
    "berkeley",
    "menlo park",
    "redwood city",
    "fremont",
    "san mateo",
}
REMOTE_TERMS = {"remote", "work from home", "wfh"}
HYBRID_TERMS = {"hybrid"}
ONSITE_TERMS = {"onsite", "on-site", "on site"}
ML_ROLE_HINTS = {
    "ml engineer",
    "machine learning engineer",
    "applied scientist",
    "ai engineer",
    "artificial intelligence engineer",
    "data scientist",
}
ML_RELEVANCE_TERMS = {
    "machine learning",
    " ml ",
    " ai ",
    "artificial intelligence",
    "data scientist",
    "applied scientist",
    "model",
    "models",
    "modeling",
    "nlp",
    "natural language processing",
    "deep learning",
    "pytorch",
    "tensorflow",
    "scikit-learn",
    "sklearn",
}
GENERIC_ROLE_WORDS = {
    "engineer",
    "scientist",
    "analyst",
    "developer",
    "manager",
    "specialist",
    "associate",
}
SENIORITY_DEALBREAKER_TERMS = {"senior", "sr", "staff", "principal", "lead"}
DEFENSE_DEALBREAKER_TERMS = {
    "defense",
    "defence",
    "military",
    "dod",
    "department of defense",
    "army",
    "navy",
    "air force",
    "marine corps",
    "space force",
    "national security",
    "homeland security",
    "government contractor",
    "federal contractor",
    "defense contractor",
    "aerospace defense",
    "weapons systems",
    "mission systems",
    "intelligence community",
    "clearance",
    "security clearance",
    "clearance required",
    "active clearance",
    "secret clearance",
    "top secret",
    "ts/sci",
    "sci clearance",
    "polygraph",
    "us citizen required",
    "u.s. citizenship required",
    "aberdeen proving ground",
}
DEFENSE_COMPANY_SIGNALS = {
    "dcs corp",
    "booz allen",
    "leidos",
    "raytheon",
    "rtx",
    "northrop grumman",
    "lockheed martin",
    "general dynamics",
    "bae systems",
    "saic",
    "caci",
    "l3harris",
    "palantir",
    "anduril",
}
JUNIOR_ROLE_TERMS = {
    "junior",
    "jr",
    "entry level",
    "entry-level",
    "new grad",
    "early career",
    "intern",
    "internship",
}
MANAGER_ROLE_TERMS = {"manager", "director", "head", "vp"}
CONTRACT_ROLE_TERMS = {
    "c2c",
    "contract",
    "contract duration",
    "contract-only",
    "contractor",
    "corp-to-corp",
    "duration",
    "freelance",
    "hourly",
    "rate open",
    "rate-open",
    "temp",
    "temporary",
    "w2 contract",
}
NON_JOB_TRAINING_ROLE_TERMS = {
    "course",
    "evaluator",
    "instructor",
    "online course",
    "search engine evaluator",
    "teacher",
    "training",
    "tutor",
}
GENERIC_SOFTWARE_TITLE_TERMS = {
    ".net developer",
    "application engineer",
    "fullstack application engineer",
    "full stack developer",
    "fullstack developer",
    "java developer",
    "software developer",
    "sql developer",
}
UNPAID_ROLE_TERMS = {
    "unpaid",
    "volunteer",
    "commission only",
    "commission-only",
    "commission based",
}
SPONSORSHIP_NEGATIVE_TERMS = {
    "no sponsorship",
    "unable to sponsor",
    "cannot sponsor",
    "without sponsorship",
    "must be authorized",
    "must be authorised",
    "authorized to work without sponsorship",
    "not sponsor",
    "must be us citizen",
    "us citizenship required",
}
SPONSORSHIP_POSITIVE_TERMS = {
    "visa",
    "sponsorship",
    "h-1b",
    "h1b",
    "visa sponsorship",
    "sponsorship available",
    "immigration",
    "research lab",
    "university",
    "large tech",
    "global company",
    "opt",
    "cpt",
}
COMPANY_PREFERENCE_TERMS = {
    "Prefer tech": {
        "bonus": {"software", "cloud", "ai", "machine learning", "data", "platform", "technology", "tech"},
        "penalty": set(),
    },
    "Prefer healthcare": {
        "bonus": {"health", "healthcare", "hospital", "clinical", "medical", "patient", "pharma", "biotech"},
        "penalty": set(),
    },
    "Prefer large companies": {
        "bonus": {"global", "enterprise", "fortune", "corporation", "inc", "technologies", "group"},
        "penalty": {"startup", "seed stage", "early stage"},
    },
    "Avoid tiny startups": {
        "bonus": set(),
        "penalty": {"startup", "seed stage", "early stage", "founding", "stealth"},
    },
    "Prefer research labs": {
        "bonus": {"research", "lab", "laboratory", "university", "institute", "scientist", "science"},
        "penalty": set(),
    },
}
YEAR_WORDS = {
    "one": 1,
    "two": 2,
    "three": 3,
    "four": 4,
    "five": 5,
    "six": 6,
    "seven": 7,
    "eight": 8,
    "nine": 9,
    "ten": 10,
}
ROLE_FAMILIES = {
    "ml_ai": {
        "target_terms": {
            "ml engineer",
            "machine learning engineer",
            "machine learning",
            "applied scientist",
            "ai engineer",
            "artificial intelligence",
            "data scientist",
        },
        "positive_terms": {
            "machine learning",
            "ml engineer",
            "ai engineer",
            "applied scientist",
            "data scientist",
            "modeling",
            "model deployment",
            "model serving",
            "nlp",
            "deep learning",
            "pytorch",
            "tensorflow",
            "scikit-learn",
            "sklearn",
            "mlops",
            "artificial intelligence",
        },
        "negative_terms": {
            "tutor",
            "expert",
            "course",
            "online course",
            "instructor",
            "teacher",
            "trainer",
            "search engine evaluator",
            "social media evaluator",
            "financial analyst",
            "marketing analyst",
            "business analyst",
            "reconciliations analyst",
            "payment processor",
            "bi developer",
            "reporting analyst",
            "accounting",
            "sales",
        },
    },
    "analytics": {
        "target_terms": {
            "data analyst",
            "bi analyst",
            "business analyst",
            "analytics engineer",
            "business intelligence",
        },
        "positive_terms": {
            "data analyst",
            "bi analyst",
            "analytics engineer",
            "business intelligence",
            "sql",
            "tableau",
            "dashboard",
            "dashboards",
            "reporting",
            "data analysis",
            "power bi",
        },
        "negative_terms": {
            "tutor",
            "course",
            "online course",
            "instructor",
            "teacher",
            "sales",
            "payment processing",
            "payment processor",
            "driver",
            "warehouse",
        },
    },
    "mlops_infra": {
        "target_terms": {
            "mlops",
            "mlops engineer",
            "ml platform",
            "ml platform engineer",
            "infrastructure engineer",
            "platform engineer",
            "senior ml engineer",
            "senior machine learning engineer",
        },
        "positive_terms": {
            "mlops",
            "ml platform",
            "ml infrastructure",
            "machine learning platform",
            "production ml",
            "production machine learning",
            "kubernetes",
            "docker",
            "spark",
            "kafka",
            "cloud",
            "aws",
            "gcp",
            "azure",
            "deployment",
            "model deployment",
            "model serving",
            "infrastructure",
            "microservices",
        },
        "negative_terms": {
            "data scientist",
            "research scientist",
            "data analyst",
            "business analyst",
            "financial analyst",
            "marketing analyst",
            "reporting analyst",
        },
    },
    "research": {
        "target_terms": {
            "research scientist",
            "applied scientist",
            "ai researcher",
            "ai research",
        },
        "positive_terms": {
            "research scientist",
            "applied scientist",
            "ai research",
            "nlp",
            "computer vision",
            "deep learning",
            "publication",
            "publications",
            "phd",
            "lab",
            "laboratory",
        },
        "negative_terms": {
            "sales",
            "marketing analyst",
            "financial analyst",
            "payment processor",
            "tutor",
        },
    },
}
JOB_ROLE_FAMILY_TERMS = {
    "ml_ai": ROLE_FAMILIES["ml_ai"]["positive_terms"],
    "analytics_bi": ROLE_FAMILIES["analytics"]["positive_terms"],
    "software": {
        "software engineer",
        "backend",
        "frontend",
        "full stack",
        "developer",
        "java",
        "python",
        "javascript",
    },
    "mlops_infra": ROLE_FAMILIES["mlops_infra"]["positive_terms"],
    "product_pm": {
        "product manager",
        "project manager",
        "program manager",
        "product owner",
        "scrum",
    },
    "business_finance_ops": {
        "business analyst",
        "financial analyst",
        "operations analyst",
        "strategy",
        "consulting",
        "supply chain",
    },
    "marketing_sales": {
        "marketing analyst",
        "growth",
        "sales",
        "account manager",
        "customer success",
    },
    "design_ux": {
        "ux",
        "ui",
        "product designer",
        "user research",
        "visual designer",
    },
    "research_science": ROLE_FAMILIES["research"]["positive_terms"],
}
US_STATES = {
    "alabama": "AL",
    "alaska": "AK",
    "arizona": "AZ",
    "arkansas": "AR",
    "california": "CA",
    "colorado": "CO",
    "connecticut": "CT",
    "delaware": "DE",
    "florida": "FL",
    "georgia": "GA",
    "hawaii": "HI",
    "idaho": "ID",
    "illinois": "IL",
    "indiana": "IN",
    "iowa": "IA",
    "kansas": "KS",
    "kentucky": "KY",
    "louisiana": "LA",
    "maine": "ME",
    "maryland": "MD",
    "massachusetts": "MA",
    "michigan": "MI",
    "minnesota": "MN",
    "mississippi": "MS",
    "missouri": "MO",
    "montana": "MT",
    "nebraska": "NE",
    "nevada": "NV",
    "new hampshire": "NH",
    "new jersey": "NJ",
    "new mexico": "NM",
    "new york": "NY",
    "north carolina": "NC",
    "north dakota": "ND",
    "ohio": "OH",
    "oklahoma": "OK",
    "oregon": "OR",
    "pennsylvania": "PA",
    "rhode island": "RI",
    "south carolina": "SC",
    "south dakota": "SD",
    "tennessee": "TN",
    "texas": "TX",
    "utah": "UT",
    "vermont": "VT",
    "virginia": "VA",
    "washington": "WA",
    "west virginia": "WV",
    "wisconsin": "WI",
    "wyoming": "WY",
}
US_STATE_ABBREVIATIONS = set(US_STATES.values())
CANADA_LOCATION_TERMS = {
    "canada",
    "alberta",
    "british columbia",
    "manitoba",
    "new brunswick",
    "newfoundland",
    "nova scotia",
    "ontario",
    "prince edward island",
    "quebec",
    "saskatchewan",
    "toronto",
    "vancouver",
    "montreal",
    "ottawa",
}
NON_US_COUNTRY_CODES = {
    "au",
    "ca",
    "de",
    "fr",
    "gb",
    "ie",
    "in",
    "it",
    "uk",
}
NON_US_LOCATION_TERMS = {
    "australia",
    "berkshire",
    "birmingham",
    "cividate al piano",
    "county dublin",
    "dublin",
    "edinburgh",
    "england",
    "eu",
    "europe",
    "european union",
    "germany",
    "india",
    "ireland",
    "italy",
    "london",
    "manchester",
    "republic of ireland",
    "scotland",
    "sheffield",
    "southampton",
    "sydney",
    "uk",
    "united kingdom",
    "uxbridge",
}
NON_US_SALARY_TERMS = {"pound", "gbp", "£"}
US_REMOTE_TERMS = {
    "remote - us",
    "remote - usa",
    "remote united states",
    "remote us",
    "remote usa",
    "remote within the us",
    "remote within us",
    "remote within united states",
    "u.s. based remote",
    "u.s. remote",
    "united states remote",
    "us based remote",
    "us remote",
    "usa based remote",
    "usa remote",
}
FEEDBACK_IGNORED_TERMS = {
    "aws",
    "cloud",
    "data",
    "engineer",
    "engineering",
    "full_time",
    "machine",
    "ml",
    "python",
    "software",
}


def split_keywords(text: str) -> list[str]:
    """Split user-entered comma/newline text into normalized keywords."""
    if not text:
        return []
    parts = re.split(r"[,;\n]+", text.lower())
    return [part.strip() for part in parts if part.strip()]


def split_target_roles(text: str) -> list[str]:
    if not text:
        return []
    parts = re.split(r"[,;/\n]+", text.lower())
    return [part.strip() for part in parts if part.strip()]


def combine_skill_inputs(manual_skills_text: str, extracted_skills: list[str]) -> str:
    combined_skills: list[str] = []
    seen_skills: set[str] = set()

    for skill in split_keywords(manual_skills_text) + extracted_skills:
        skill_key = skill.lower()
        if skill_key not in seen_skills:
            combined_skills.append(skill)
            seen_skills.add(skill_key)

    return ", ".join(combined_skills)


def format_profile_signals_for_display(skills_text: str) -> str:
    """Clean skill/profile signals for display without changing matching inputs."""
    if not skills_text:
        return ""

    noise_terms = {"certified", "some"}
    display_skills: list[str] = []
    seen: set[str] = set()
    for raw_skill in re.split(r"[,;\n.]+", skills_text):
        words = [
            word
            for word in re.sub(r"[^a-zA-Z0-9+#/ -]+", " ", raw_skill).split()
            if word.lower() not in noise_terms
        ]
        cleaned_skill = " ".join(words).strip(" -")
        normalized = cleaned_skill.lower()
        if not normalized or normalized in noise_terms or normalized in seen:
            continue
        display_skills.append(format_skill_name(cleaned_skill))
        seen.add(normalized)
    return ", ".join(display_skills)


def parse_json_like_value(value: object) -> object:
    if not isinstance(value, str):
        return value
    text = value.strip()
    if not (text.startswith("{") and text.endswith("}")):
        return value
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        try:
            return ast.literal_eval(text)
        except (ValueError, SyntaxError):
            return value


def extract_numbers(text: str) -> list[int]:
    numbers: list[int] = []
    for match in re.findall(r"\d[\d,]*(?:\.\d+)?", text):
        number = int(float(match.replace(",", "")))
        if number >= 1_000:
            numbers.append(number)
    return numbers


def parse_number(value: object) -> int:
    if value in (None, "") or pd.isna(value):
        return 0
    try:
        return int(float(str(value).replace(",", "")))
    except ValueError:
        numbers = extract_numbers(str(value))
        return max(numbers) if numbers else 0


def parse_salary(value: str) -> int:
    """Convert a salary value into a usable scoring number when possible."""
    if pd.isna(value):
        return 0

    parsed_value = parse_json_like_value(value)
    if isinstance(parsed_value, dict):
        for key in ["value", "minValue", "maxValue"]:
            number = parse_number(parsed_value.get(key))
            if number:
                return number

    numbers = extract_numbers(str(value))
    return max(numbers) if numbers else 0


def make_salary_text(value: object) -> str:
    if pd.isna(value):
        return "Unknown"

    parsed_value = parse_json_like_value(value)
    if isinstance(parsed_value, dict):
        text = parsed_value.get("text")
        if text:
            return str(text).strip()
        number = parse_salary(str(parsed_value))
        return f"{number:,}" if number else "Unknown"

    text = str(value).strip()
    return text if text else "Unknown"


def normalize_location_text(text: str) -> str:
    return re.sub(r"[^a-z0-9]+", " ", text.lower()).strip()


def normalize_search_text(text: str) -> str:
    return f" {re.sub(r'[^a-z0-9+#.]+', ' ', text.lower()).strip()} "


def contains_phrase(text: str, phrase: str) -> bool:
    return normalize_search_text(phrase).strip() in normalize_search_text(text)


def score_target_role_relevance(
    target_role: str,
    title: str,
    description: str,
) -> tuple[int, str]:
    roles = split_target_roles(target_role)
    if not roles:
        return 0, ""

    title_text = normalize_search_text(title)
    searchable_text = normalize_search_text(f"{title} {description}")
    points = 0
    reasons: list[str] = []

    title_matches = [role for role in roles if role in title_text]
    if title_matches:
        points += 40
        reasons.append(
            f"Target role relevance: +40 title match: {', '.join(title_matches)}"
        )
    else:
        partial_matches = [
            role
            for role in roles
            if any(
                word in title_text
                for word in role.split()
                if len(word) > 2 and word not in GENERIC_ROLE_WORDS
            )
        ]
        if partial_matches:
            points += 20
            reasons.append(
                "Target role relevance: "
                f"+20 related title terms: {', '.join(partial_matches)}"
            )
        else:
            reasons.append("Target role relevance: +0 not directly found in title")

    if any(role in ML_ROLE_HINTS or "machine learning" in role for role in roles):
        matched_ml_terms = [
            term.strip()
            for term in ML_RELEVANCE_TERMS
            if term in searchable_text
        ]
        if matched_ml_terms:
            bonus = min(35, 5 * len(matched_ml_terms))
            points += bonus
            reasons.append(
                "Target role relevance: "
                f"+{bonus} ML/applied science terms: {', '.join(matched_ml_terms[:6])}"
            )
        elif not title_matches:
            points -= 25
            reasons.append(
                "Target role relevance: -25 no ML/applied science evidence found"
            )

    return points, "; ".join(reasons)


def preferred_location_terms(preferred_location: str) -> set[str]:
    normalized = normalize_location_text(preferred_location)
    terms = set(normalized.split())
    for state_name, abbreviation in US_STATES.items():
        if state_name in normalized:
            terms.add(abbreviation.lower())
    return terms


def is_us_preference(preferred_location: str) -> bool:
    raw_preference = preferred_location.lower()
    normalized = normalize_location_text(preferred_location)
    words = set(normalized.split())
    if "us only" in raw_preference or "u.s. only" in raw_preference:
        return True
    if "nyc" in words or "new york city" in normalized:
        return True
    if normalized in {"us", "usa", "u s", "united states"}:
        return True
    if words & {"us", "usa"}:
        return True
    if any(state_name in normalized for state_name in US_STATES):
        return True
    return bool({word.upper() for word in words} & US_STATE_ABBREVIATIONS)


def location_looks_us(location: str) -> bool:
    normalized = normalize_location_text(location)
    if location_looks_non_us(location):
        return False
    words = set(normalized.split())
    if words & {"us", "usa"} or "united states" in normalized:
        return True
    if any(state_name in normalized for state_name in US_STATES):
        return True
    return bool({word.upper() for word in words} & US_STATE_ABBREVIATIONS)


def location_looks_canadian(location: str) -> bool:
    normalized = normalize_location_text(location)
    return any(term in normalized for term in CANADA_LOCATION_TERMS)


def location_looks_non_us(location: str) -> bool:
    raw_location = str(location).lower()
    normalized = normalize_location_text(location)
    words = normalized.split()
    if set(words) & {"us", "usa"} or "united states" in normalized:
        return False
    if "£" in raw_location or any(term in normalized for term in NON_US_SALARY_TERMS - {"£"}):
        return True
    if any(term in normalized for term in NON_US_LOCATION_TERMS | CANADA_LOCATION_TERMS):
        return True

    raw_parts = [
        normalize_location_text(part)
        for part in re.split(r"[,/|]", str(location))
        if normalize_location_text(part)
    ]
    for part in raw_parts:
        part_words = part.split()
        if len(part_words) == 1 and part_words[0] in NON_US_COUNTRY_CODES:
            return True

    if words and words[-1] in NON_US_COUNTRY_CODES:
        return True
    return False


def location_has_explicit_non_us_signal(location: str) -> bool:
    raw_text = str(location)
    normalized = normalize_location_text(raw_text)
    words = normalized.split()
    has_us_context = (
        set(words) & {"us", "usa"}
        or "united states" in normalized
        or bool({word.upper() for word in words} & US_STATE_ABBREVIATIONS)
        or any(state_name in normalized for state_name in US_STATES)
    )

    if "£" in raw_text.lower() or any(term in normalized for term in NON_US_SALARY_TERMS - {"£"}):
        return True
    if any(term in normalized for term in NON_US_LOCATION_TERMS | CANADA_LOCATION_TERMS):
        return not has_us_context

    raw_parts = [
        normalize_location_text(part)
        for part in re.split(r"[,/|]", raw_text)
        if normalize_location_text(part)
    ]
    for part in raw_parts:
        part_words = part.split()
        if len(part_words) != 1:
            continue
        code = part_words[0]
        if code in {"ca", "in"}:
            return not has_us_context
        if code in NON_US_COUNTRY_CODES:
            return True

    if words:
        last_word = words[-1]
        if last_word in {"ca", "in"}:
            return not has_us_context
        if last_word in NON_US_COUNTRY_CODES:
            return True
    return False


def location_matches_bay_area(location_context: str) -> bool:
    normalized = normalize_location_text(location_context)
    return any(term in normalized for term in BAY_AREA_TERMS)


def location_matches_remote(location_context: str) -> bool:
    normalized = normalize_location_text(location_context)
    return (
        "remote" in normalized.split()
        or "work from home" in normalized
        or "wfh" in normalized.split()
    )


def location_has_us_remote_signal(location_context: str) -> bool:
    normalized = normalize_search_text(location_context)
    return any(
        normalize_search_text(term).strip() in normalized
        for term in US_REMOTE_TERMS
    )


def find_location_terms(text: str, terms: set[str]) -> list[str]:
    normalized = normalize_search_text(text)
    matches: list[str] = []
    for term in terms:
        normalized_term = normalize_search_text(term).strip()
        if re.search(rf"(?<![a-z0-9]){re.escape(normalized_term)}(?![a-z0-9])", normalized):
            matches.append(term)
    return sorted(matches)


def infer_location_source(row: pd.Series) -> str:
    source = str(row.get("location_source", "")).strip()
    if source:
        return source
    location = str(row.get("location", "")).strip()
    if not location or location.lower() == "unknown":
        return "unknown"
    return "ambiguous_location"


def location_confidence_from_source(location_source: str, location: str) -> str:
    if not location or location.lower() == "unknown" or location_source == "unknown":
        return "low"
    if location_source == "job_posting_location":
        return "high"
    if location_source in {"raw_location_field", "description_inferred"}:
        return "medium"
    if location_source in {"company_location_if_only_available", "ambiguous_location"}:
        return "low"
    return "medium"


def extract_location_signals(job_row: pd.Series) -> dict[str, object]:
    title = str(job_row.get("title", ""))
    location = str(job_row.get("location", ""))
    raw_location = str(job_row.get("raw_location", location))
    salary_text = str(job_row.get("salary_text", job_row.get("salary", "")))
    description = str(job_row.get("description", ""))
    location_source = infer_location_source(job_row)
    location_text = f"{title} {location} {raw_location} {description}"

    location_only_text = f"{location} {raw_location} {salary_text}"
    remote_matches = find_location_terms(location_text, REMOTE_TERMS)
    explicit_remote_location_matches = find_location_terms(location_only_text, REMOTE_TERMS)
    hybrid_matches = find_location_terms(location_text, HYBRID_TERMS)
    onsite_matches = find_location_terms(location_text, ONSITE_TERMS)
    bay_area_matches = find_location_terms(location_text, BAY_AREA_TERMS)
    normalized_location = normalize_search_text(f"{location} {raw_location}")
    us_matches = []
    if location_looks_us(f"{location} {raw_location}"):
        us_matches.append("US/location state signal")
    california_matches = []
    if (
        " california " in normalize_search_text(location_text)
        or " ca us " in normalized_location
        or " ca usa " in normalized_location
    ):
        california_matches.append("California/CA")

    is_remote = bool(remote_matches)
    is_us_remote = location_has_us_remote_signal(location_text)
    is_hybrid = bool(hybrid_matches)
    is_bay_area = bool(bay_area_matches)
    is_california = bool(california_matches) or is_bay_area
    has_non_us_signal = location_looks_non_us(location_only_text)
    has_explicit_non_us_signal = location_has_explicit_non_us_signal(location_only_text)
    is_us = bool(us_matches) and not has_non_us_signal
    normalized_location_only = normalize_location_text(f"{location} {raw_location}")
    remote_only_location = normalized_location_only in {
        "remote",
        "remote remote",
        "anywhere",
        "worldwide",
    }
    is_unknown = not location.strip() or location.strip().lower() == "unknown"
    is_non_us = (
        not is_unknown
        and not is_us
        and not remote_only_location
        and (has_explicit_non_us_signal or has_non_us_signal or not is_remote)
    )

    matched_phrases = {
        "remote": remote_matches,
        "hybrid": hybrid_matches,
        "onsite": onsite_matches,
        "us": us_matches,
        "non_us": ["non-US country/location signal"] if has_non_us_signal else [],
        "california": california_matches,
        "bay_area": bay_area_matches,
    }

    return {
        "location": location,
        "raw_location": raw_location,
        "location_source": location_source,
        "location_confidence": location_confidence_from_source(location_source, location),
        "is_us_location": is_us,
        "is_remote": is_remote,
        "is_us_remote_location": is_us_remote,
        "has_explicit_non_us_location_signal": has_explicit_non_us_signal,
        "is_explicit_remote_location": bool(explicit_remote_location_matches),
        "is_hybrid": is_hybrid,
        "is_california": is_california,
        "is_bay_area": is_bay_area,
        "is_non_us_location": is_non_us,
        "location_matched_phrases": matched_phrases,
    }


def evaluate_location_preference(
    preferred_location: str,
    location_signals: dict[str, object],
    strictness: str = "Balanced",
) -> tuple[bool, str]:
    if not preferred_location:
        return True, "No location preference entered."

    normalized_preference = normalize_location_text(preferred_location)
    wants_us = is_us_preference(preferred_location)
    wants_remote = "remote" in normalized_preference
    wants_bay_area = "bay area" in normalized_preference
    wants_california = "california" in normalized_preference or " ca " in f" {normalized_preference} "

    matched = False
    reasons: list[str] = []
    if wants_us:
        matched = matched or bool(location_signals["is_us_location"])
        reasons.append("US preference")
    if wants_remote:
        matched = matched or bool(location_signals["is_remote"])
        reasons.append("remote preference")
    if wants_bay_area:
        matched = matched or bool(location_signals["is_bay_area"])
        reasons.append("Bay Area preference")
    if wants_california:
        matched = matched or bool(location_signals["is_california"])
        reasons.append("California preference")
    if not (wants_us or wants_remote or wants_bay_area or wants_california):
        matched = normalized_preference in normalize_location_text(
            f"{location_signals.get('raw_location', '')} {location_signals.get('location', '')}"
        )
        reasons.append("text location preference")

    unknown_location = (
        location_signals.get("location_source") == "unknown"
        and not location_signals["is_non_us_location"]
    )
    if (
        strictness == "Strict"
        and wants_us
        and location_signals["is_non_us_location"]
    ):
        if location_signals.get("is_us_remote_location", False):
            return True, "PASS: strict US preference allowed explicit US-remote location."
        return False, f"FAIL: strict location filter excludes clear non-US location for {' / '.join(reasons)}."
    if matched:
        return True, f"PASS: matched {' / '.join(reasons)}."
    if (
        strictness == "Strict"
        and wants_us
        and location_signals.get("is_us_remote_location", False)
    ):
        return True, "PASS: strict US preference allowed explicit US-remote location."
    if unknown_location:
        return True, "PASS: location unknown, not failed automatically."
    if strictness == "Strict" and location_signals["is_non_us_location"]:
        return False, f"FAIL: strict location filter excludes clear non-US location for {' / '.join(reasons)}."
    if strictness == "Flexible":
        return True, "PASS: Flexible location strictness allows broader exploration."
    return False, f"FAIL: did not match {' / '.join(reasons)}."


def score_location_signals(
    preferred_location: str,
    location_signals: dict[str, object],
    strictness: str = "Balanced",
) -> tuple[int, str]:
    if not preferred_location:
        return 0, ""

    passed, reason = evaluate_location_preference(
        preferred_location,
        location_signals,
        strictness=strictness,
    )
    normalized_preference = normalize_location_text(preferred_location)

    if "remote" in normalized_preference and location_signals["is_remote"]:
        return 25, f"Location preference: +25 remote match. {reason}"
    if "bay area" in normalized_preference and location_signals["is_bay_area"]:
        return 25, f"Location preference: +25 Bay Area match. {reason}"
    if "california" in normalized_preference and location_signals["is_california"]:
        return 20, f"Location preference: +20 California match. {reason}"
    if is_us_preference(preferred_location) and location_signals["is_us_location"]:
        return 15, f"Location preference: +15 US match. {reason}"
    if passed and location_signals["location_confidence"] == "low":
        return 0, f"Location preference: +0 low-confidence/unknown location. {reason}"
    if passed:
        return 10, f"Location preference: +10 location allowed. {reason}"
    if strictness == "Strict":
        return -45, f"Location preference: -45 strict location mismatch. {reason}"
    return -20, f"Location preference: -20 location mismatch. {reason}"


def score_location(
    preferred_location: str,
    location: str,
    location_context: str = "",
) -> tuple[int, str]:
    if not preferred_location:
        return 0, ""

    raw_preference = preferred_location.lower()
    normalized_preference = normalize_location_text(preferred_location)
    normalized_location = normalize_location_text(location)
    normalized_context = normalize_location_text(f"{location} {location_context}")

    if not normalized_location or normalized_location == "unknown":
        return -5, "Location preference: -5 location unknown"

    if "remote" in normalized_preference and location_matches_remote(normalized_context):
        return 25, "Location preference: +25 remote match"

    if "bay area" in normalized_preference and location_matches_bay_area(normalized_context):
        return 25, "Location preference: +25 Bay Area match"

    if "us only" in raw_preference or "u.s. only" in raw_preference:
        if location_looks_us(location):
            return 10, "Location preference: +10 US-only preference satisfied"
        return -35, "Location preference: -35 US-only preference but job appears non-US"

    if is_us_preference(preferred_location) and not location_looks_us(location):
        return -35, "Location preference: -35 US preference but job appears non-US"

    terms = preferred_location_terms(preferred_location)
    location_words = set(normalized_location.split())
    if normalized_preference and normalized_preference in normalized_location:
        return 20, "Location preference: +20 preferred location match"
    if terms and terms & location_words:
        return 20, "Location preference: +20 preferred location/state match"

    if "remote" in normalized_preference or "bay area" in normalized_preference:
        return -15, "Location preference: -15 not remote or Bay Area"

    return -10, "Location preference: -10 differs from preference"


def score_dealbreakers(
    dealbreakers: list[str],
    title: str,
    description: str,
) -> tuple[int, str]:
    if not dealbreakers:
        return 0, ""

    dealbreakers = [dealbreaker.lower().strip() for dealbreaker in dealbreakers]
    title_text = normalize_search_text(title)
    description_text = normalize_search_text(description)
    searchable_text = f"{title_text} {description_text}"
    penalty = 0
    reasons: list[str] = []

    if any(term in {"senior", "sr", "staff"} for term in dealbreakers):
        matched_terms = [
            term
            for term in SENIORITY_DEALBREAKER_TERMS
            if re.search(rf"\b{re.escape(term)}\.?\b", title.lower())
        ]
        if matched_terms:
            penalty -= 60
            reasons.append(
                f"Dealbreaker penalty: -60 seniority in title: {', '.join(matched_terms)}"
            )

    if any(term in {"defense", "military"} for term in dealbreakers):
        matched_terms = [
            term
            for term in DEFENSE_DEALBREAKER_TERMS
            if normalize_search_text(term).strip() in searchable_text
        ]
        if matched_terms:
            penalty -= 60
            reasons.append(
                f"Dealbreaker penalty: -60 defense/military terms: {', '.join(matched_terms)}"
            )

    if any(term in {"5+ years", "5 years", "five years"} for term in dealbreakers):
        experience_patterns = [
            r"\b[5-9]\+?\s*years\b",
            r"\bfive\s+years\b",
            r"\bsix\s+years\b",
            r"\bseven\s+years\b",
        ]
        if any(re.search(pattern, description.lower()) for pattern in experience_patterns):
            penalty -= 35
            reasons.append("Dealbreaker penalty: -35 5+ years or more mentioned")

    generic_matches = [
        dealbreaker
        for dealbreaker in dealbreakers
        if dealbreaker not in {"senior", "sr", "staff", "defense", "military", "5+ years", "5 years", "five years"}
        and dealbreaker in searchable_text
    ]
    if generic_matches:
        generic_penalty = 25 * len(generic_matches)
        penalty -= generic_penalty
        reasons.append(
            f"Dealbreaker penalty: -{generic_penalty} found: {', '.join(generic_matches)}"
        )

    return penalty, "; ".join(reasons)


def extract_required_years(text: str) -> int:
    years, _ = extract_years_requirement(text)
    return years or 0


def extract_years_requirement(text: str) -> tuple[int | None, str]:
    text = text.lower()
    patterns = [
        r"\b(?:minimum|at least)\s+(\d{1,2})\s*\+?\s*(?:years|yrs)\b(?:\s+of\s+experience)?",
        r"\b(\d{1,2})\s*-\s*\d{1,2}\s*(?:years|yrs)\b(?:\s+of\s+experience)?",
        r"\b(\d{1,2})\s*\+?\s*(?:years|yrs)\b(?:\s+of\s+experience)?",
        r"\b(\d{1,2})\s+or\s+more\s+(?:years|yrs)\b(?:\s+of\s+experience)?",
    ]
    matches: list[tuple[int, str]] = []
    for pattern in patterns:
        for match in re.finditer(pattern, text):
            years = int(match.group(1))
            phrase = match.group(0)
            nearby_text = text[match.start() : match.end() + 12]
            if (
                years <= 15
                and not re.search(r"\b(?:salary|pay|paid|old|age)\b", phrase)
                and "ago" not in nearby_text
            ):
                matches.append((years, phrase))

    for word, value in YEAR_WORDS.items():
        match = re.search(
            rf"\b(?:minimum|at least)?\s*{word}\s+(?:years|yrs)\b(?:\s+of\s+experience)?",
            text,
        )
        if match and "ago" not in text[match.start() : match.end() + 12]:
            matches.append((value, match.group(0)))

    if not matches:
        return None, "Unknown"
    return min(matches, key=lambda item: item[0])


def find_terms(text: str, terms: set[str]) -> list[str]:
    normalized = normalize_search_text(text)
    matches: list[str] = []
    for term in terms:
        normalized_term = normalize_search_text(term).strip()
        if re.search(rf"(?<![a-z0-9]){re.escape(normalized_term)}(?![a-z0-9])", normalized):
            matches.append(term)
    return sorted(matches)


def infer_employment_type(
    is_contract_or_temp: bool,
    is_unpaid_or_commission: bool,
    text: str,
) -> str:
    normalized = normalize_search_text(text)
    if is_unpaid_or_commission:
        return "unpaid_or_commission"
    if is_contract_or_temp:
        return "contract_or_temp"
    if " internship " in normalized or " intern " in normalized:
        return "internship"
    if " part time " in normalized:
        return "part_time"
    if " full time " in normalized:
        return "full_time"
    return "Unknown"


def infer_job_role_family(text: str) -> str:
    best_family = "unknown"
    best_count = 0
    for family, terms in JOB_ROLE_FAMILY_TERMS.items():
        count = len(find_terms(text, terms))
        if count > best_count:
            best_family = family
            best_count = count
    return best_family


def find_defense_military_matches(
    company: str,
    metadata_text: str,
) -> list[str]:
    keyword_matches = find_terms(metadata_text, DEFENSE_DEALBREAKER_TERMS)
    company_matches = find_terms(company, DEFENSE_COMPANY_SIGNALS)
    metadata_company_matches = find_terms(metadata_text, DEFENSE_COMPANY_SIGNALS)

    matches: list[str] = []
    seen_matches: set[str] = set()
    for match in company_matches + metadata_company_matches + keyword_matches:
        match_key = match.lower()
        if match_key not in seen_matches:
            matches.append(match)
            seen_matches.add(match_key)
    return matches


def bucket_employee_count(employee_count: float) -> str:
    if employee_count < 50:
        return "tiny"
    if employee_count < 100:
        return "small"
    if employee_count < 1_000:
        return "medium"
    if employee_count < 10_000:
        return "large"
    return "enterprise"


def extract_employee_count(row: pd.Series) -> tuple[float | None, str]:
    count_columns = [
        "employee_count",
        "employees",
        "num_employees",
        "company_employee_count",
        "company_size",
    ]
    for column in count_columns:
        value = row.get(column)
        parsed_value = safe_numeric_or_none(value)
        if parsed_value is not None and parsed_value > 0:
            return parsed_value, str(value)
    return None, ""


def infer_company_size(row: pd.Series, metadata_text: str) -> dict[str, object]:
    employee_count, employee_text = extract_employee_count(row)
    if employee_count is not None:
        size = bucket_employee_count(employee_count)
        return {
            "company_size_inferred": size,
            "company_size_confidence": "high",
            "company_size_match_text": employee_text,
        }

    tiny_matches = find_terms(metadata_text, TINY_COMPANY_SIGNALS)
    large_matches = find_terms(metadata_text, LARGE_COMPANY_SIGNALS)
    normalized_text = normalize_search_text(metadata_text)

    numeric_size_patterns = [
        (r"\b(?:under|fewer than|less than)\s+50\s+employees\b", "tiny"),
        (r"\b(?:under|fewer than|less than)\s+100\s+employees\b", "small"),
        (r"\b(\d{2,6})\+?\s+employees\b", ""),
    ]
    for pattern, forced_size in numeric_size_patterns:
        match = re.search(pattern, normalized_text)
        if not match:
            continue
        if forced_size:
            return {
                "company_size_inferred": forced_size,
                "company_size_confidence": "low",
                "company_size_match_text": match.group(0).strip(),
            }
        employee_count = safe_numeric_or_none(match.group(1))
        if employee_count is not None:
            return {
                "company_size_inferred": bucket_employee_count(employee_count),
                "company_size_confidence": "medium",
                "company_size_match_text": match.group(0).strip(),
            }

    if tiny_matches:
        size = "tiny" if any("50" in match or "startup" in match or "seed" in match for match in tiny_matches) else "small"
        return {
            "company_size_inferred": size,
            "company_size_confidence": "low",
            "company_size_match_text": ", ".join(tiny_matches),
        }
    if large_matches:
        size = "enterprise" if any(match in {"fortune 500", "multinational", "worldwide", "thousands of employees"} for match in large_matches) else "large"
        return {
            "company_size_inferred": size,
            "company_size_confidence": "low",
            "company_size_match_text": ", ".join(large_matches),
        }

    return {
        "company_size_inferred": "Unknown",
        "company_size_confidence": "low",
        "company_size_match_text": "",
    }


def extract_job_constraints(job_row: pd.Series) -> dict[str, object]:
    title = str(job_row.get("title", ""))
    company = str(job_row.get("company", ""))
    location = str(job_row.get("location", ""))
    raw_location = str(job_row.get("raw_location", location))
    salary_text = str(job_row.get("salary_text", job_row.get("salary", "")))
    description = str(job_row.get("description", ""))
    employment_type = str(job_row.get("employment_type", ""))
    metadata_text = " ".join(
        clean
        for clean in [
            title,
            company,
            location,
            raw_location,
            salary_text,
            description,
            employment_type,
        ]
        if clean
    )
    title_text = title

    required_years_min, years_text = extract_years_requirement(description)
    contract_matches = find_terms(metadata_text, CONTRACT_ROLE_TERMS)
    unpaid_matches = find_terms(metadata_text, UNPAID_ROLE_TERMS)
    senior_matches = [
        term
        for term in SENIORITY_DEALBREAKER_TERMS
        if re.search(rf"\b{re.escape(term)}\.?\b", title.lower())
    ]
    junior_matches = find_terms(title_text, JUNIOR_ROLE_TERMS)
    manager_matches = find_terms(title_text, MANAGER_ROLE_TERMS)
    defense_matches = find_defense_military_matches(company, metadata_text)
    sponsorship_positive_matches = find_terms(metadata_text, SPONSORSHIP_POSITIVE_TERMS)
    sponsorship_negative_matches = find_terms(metadata_text, SPONSORSHIP_NEGATIVE_TERMS)
    company_size_signals = infer_company_size(job_row, metadata_text)
    company_size = str(company_size_signals["company_size_inferred"]).lower()

    location_signals = extract_location_signals(job_row)
    is_contract_or_temp = bool(contract_matches)
    is_unpaid_or_commission = bool(unpaid_matches)

    matched_phrases = {
        "years_requirement": [] if years_text == "Unknown" else [years_text],
        "contract_or_temp": contract_matches,
        "unpaid_or_commission": unpaid_matches,
        "senior_staff_principal": senior_matches,
        "junior_entry_level": junior_matches,
        "manager_director": manager_matches,
        "defense_military_clearance": defense_matches,
        "visa_sponsorship": sponsorship_positive_matches,
        "no_sponsorship": sponsorship_negative_matches,
    }

    return {
        "required_years_min": required_years_min,
        "years_requirement_text": years_text,
        "salary_number": job_row.get("salary_number", 0),
        "employment_type_inferred": infer_employment_type(
            is_contract_or_temp,
            is_unpaid_or_commission,
            metadata_text,
        ),
        "is_contract_or_temp": is_contract_or_temp,
        "is_unpaid_or_commission": is_unpaid_or_commission,
        "contract_temp_text": ", ".join(contract_matches) if contract_matches else "",
        "unpaid_commission_text": ", ".join(unpaid_matches) if unpaid_matches else "",
        "is_senior_staff_principal": bool(senior_matches),
        "is_junior_entry_level": bool(junior_matches),
        "is_manager_director": bool(manager_matches),
        "is_defense_military_clearance": bool(defense_matches),
        "defense_match_text": ", ".join(defense_matches) if defense_matches else "",
        "company_size_inferred": company_size_signals["company_size_inferred"],
        "company_size_confidence": company_size_signals["company_size_confidence"],
        "company_size_match_text": company_size_signals["company_size_match_text"],
        "is_tiny_startup": company_size == "tiny",
        "is_small_company": company_size == "small",
        "is_medium_company": company_size == "medium",
        "is_large_company": company_size == "large",
        "is_enterprise_company": company_size == "enterprise",
        "has_visa_sponsorship_signal": bool(sponsorship_positive_matches),
        "has_no_sponsorship_signal": bool(sponsorship_negative_matches),
        "raw_location": location_signals["raw_location"],
        "location_source": location_signals["location_source"],
        "location_confidence": location_signals["location_confidence"],
        "is_us_location": location_signals["is_us_location"],
        "is_remote": location_signals["is_remote"],
        "is_us_remote_location": location_signals["is_us_remote_location"],
        "has_explicit_non_us_location_signal": location_signals[
            "has_explicit_non_us_location_signal"
        ],
        "is_hybrid": location_signals["is_hybrid"],
        "is_california": location_signals["is_california"],
        "is_bay_area": location_signals["is_bay_area"],
        "is_california_or_bay_area": bool(location_signals["is_california"])
        or bool(location_signals["is_bay_area"]),
        "is_non_us_location": location_signals["is_non_us_location"],
        "location_matched_phrases": location_signals["location_matched_phrases"],
        "role_family_inferred": infer_job_role_family(metadata_text),
        "constraint_matched_phrases": matched_phrases,
    }


def avoid_contract_temp_active(advanced_preferences: dict) -> bool:
    return bool(
        advanced_preferences.get("avoid_contract_roles", False)
        or advanced_preferences.get("avoid_contract_temp", False)
    )


def avoid_unpaid_commission_active(advanced_preferences: dict) -> bool:
    return bool(
        advanced_preferences.get("avoid_unpaid_roles", False)
        or advanced_preferences.get("avoid_unpaid_commission", False)
    )


def avoid_internship_active(advanced_preferences: dict) -> bool:
    return bool(
        advanced_preferences.get("avoid_internship_roles", False)
        or advanced_preferences.get("seniority_preference") in {
            "Senior+ preferred",
            "Exclude Junior/Entry-level",
        }
    )


def derive_seniority_filter_flags(seniority_preference: str) -> dict[str, bool]:
    avoid_senior = seniority_preference in {
        "Entry-level / New Grad only",
        "Junior to Mid-level",
        "Exclude Senior/Staff/Principal",
    }
    avoid_junior = seniority_preference in {
        "Mid to Senior allowed",
        "Senior+ preferred",
        "Exclude Junior/Entry-level",
    }
    avoid_manager = seniority_preference in {
        "Entry-level / New Grad only",
        "Junior to Mid-level",
        "Exclude Senior/Staff/Principal",
    }
    return {
        "avoid_senior_roles": avoid_senior,
        "avoid_junior_roles": avoid_junior,
        "avoid_manager_roles": avoid_manager,
    }


def score_seniority_preference(
    constraints: dict[str, object],
    seniority_preference: str,
    title: str,
) -> tuple[int, str]:
    if seniority_preference == "No preference":
        return 0, ""

    is_senior = bool(constraints.get("is_senior_staff_principal"))
    is_junior = bool(constraints.get("is_junior_entry_level"))
    is_manager = bool(constraints.get("is_manager_director"))
    normalized_title = normalize_search_text(title)
    has_entry_signals = (
        is_junior
        or " analyst " in normalized_title
        or " associate " in normalized_title
        or " early career " in normalized_title
        or re.search(r"\b(?:i|ii|1|2)\b", title.lower()) is not None
    )

    if seniority_preference == "Entry-level / New Grad only":
        if is_senior or is_manager:
            return -90, "Seniority preference: -90 excludes senior/manager roles"
        if has_entry_signals:
            return 25, "Seniority preference: +25 entry-level/new grad signal"
        return -15, "Seniority preference: -15 no clear entry-level signal"

    if seniority_preference == "Junior to Mid-level":
        if is_senior or is_manager:
            return -70, "Seniority preference: -70 senior/manager role"
        if has_entry_signals:
            return 15, "Seniority preference: +15 junior/mid-level signal"
        return 5, "Seniority preference: +5 no seniority conflict"

    if seniority_preference == "Mid to Senior allowed":
        if is_junior:
            return -80, "Seniority preference: -80 junior/entry-level role"
        return 5, "Seniority preference: +5 mid/senior-compatible role"

    if seniority_preference == "Senior+ preferred":
        if is_senior:
            return 25, "Seniority preference: +25 senior/staff/principal signal"
        if is_junior:
            return -45, "Seniority preference: -45 junior/entry-level role"
        return -5, "Seniority preference: -5 no senior+ signal"

    if seniority_preference == "Exclude Senior/Staff/Principal":
        if is_senior or is_manager:
            return -120, "Seniority preference: -120 excluded senior/manager role"
        return 0, ""

    if seniority_preference == "Exclude Junior/Entry-level":
        if is_junior:
            return -100, "Seniority preference: -100 excluded junior/entry-level role"
        return 0, ""

    return 0, ""


def safe_numeric_or_none(value: object) -> float | None:
    if value is None:
        return None
    try:
        if pd.isna(value):
            return None
    except TypeError:
        return None

    text = str(value).strip()
    if not text or text.lower() == "unknown":
        return None

    try:
        return float(text)
    except ValueError:
        return None


def existing_columns(df: pd.DataFrame, columns: list[str]) -> list[str]:
    return [column for column in columns if column in df.columns]


def evaluate_active_constraint_checks(
    constraints: dict[str, object],
    advanced_preferences: dict,
) -> list[dict[str, str]]:
    checks: list[dict[str, str]] = []

    def add_check(name: str, active: bool, passed: bool, detail: str) -> None:
        if active:
            checks.append(
                {
                    "filter": name,
                    "status": "PASS" if passed else "FAIL",
                    "detail": detail,
                }
            )

    max_years = advanced_preferences.get("max_required_years", "No limit")
    required_years = constraints.get("required_years_min")
    max_allowed = safe_numeric_or_none(advanced_preferences.get("max_required_years_value"))
    if max_allowed is None:
        max_allowed = safe_numeric_or_none(str(max_years).split()[0])
    if max_years != "No limit" and max_allowed is not None:
        required_years_value = safe_numeric_or_none(required_years)
        passed = required_years_value is None or required_years_value <= max_allowed
        add_check(
            "Maximum required years",
            True,
            passed,
            f"required={required_years_value if required_years_value is not None else 'Unknown'}, max={max_allowed}",
        )

    minimum_salary = safe_numeric_or_none(advanced_preferences.get("minimum_salary"))
    salary_strictness = advanced_preferences.get("salary_strictness", "Soft")
    salary_number = safe_numeric_or_none(constraints.get("salary_number"))
    if minimum_salary and minimum_salary > 0:
        salary_unknown = salary_number is None or salary_number <= 0
        passed = salary_strictness != "Strict" or salary_unknown or salary_number >= minimum_salary
        add_check(
            "Minimum salary",
            True,
            passed,
            f"salary={salary_number if not salary_unknown else 'Unknown'}, min={minimum_salary}, strictness={salary_strictness}",
        )

    add_check(
        "Avoid contract/temp",
        avoid_contract_temp_active(advanced_preferences),
        not bool(constraints.get("is_contract_or_temp")),
        str(constraints.get("constraint_matched_phrases", {}).get("contract_or_temp", [])),
    )
    add_check(
        "Avoid unpaid/commission-only",
        avoid_unpaid_commission_active(advanced_preferences),
        not bool(constraints.get("is_unpaid_or_commission")),
        str(constraints.get("constraint_matched_phrases", {}).get("unpaid_or_commission", [])),
    )
    add_check(
        "Avoid senior/staff/principal",
        advanced_preferences.get("avoid_senior_roles", False),
        not bool(constraints.get("is_senior_staff_principal")),
        str(constraints.get("constraint_matched_phrases", {}).get("senior_staff_principal", [])),
    )
    add_check(
        "Avoid junior",
        advanced_preferences.get("avoid_junior_roles", False),
        not bool(constraints.get("is_junior_entry_level")),
        str(constraints.get("constraint_matched_phrases", {}).get("junior_entry_level", [])),
    )
    add_check(
        "Avoid internship roles",
        avoid_internship_active(advanced_preferences),
        str(constraints.get("employment_type_inferred", "")).lower() != "internship",
        str(constraints.get("employment_type_inferred", "")),
    )
    add_check(
        "Avoid manager/director",
        advanced_preferences.get("avoid_manager_roles", False),
        not bool(constraints.get("is_manager_director")),
        str(constraints.get("constraint_matched_phrases", {}).get("manager_director", [])),
    )
    add_check(
        "Strict location",
        advanced_preferences.get("location_strictness", "Balanced") == "Strict",
        bool(constraints.get("passes_location_filter", True)),
        str(constraints.get("location_filter_reason", "")),
    )
    add_check(
        "Strict role fit",
        advanced_preferences.get("role_fit_strictness", "Balanced") == "Strict",
        bool(constraints.get("passes_role_fit_filter", True)),
        f"role_fit_score={constraints.get('role_fit_score', 'Unknown')}",
    )
    add_check(
        "Avoid defense/military",
        advanced_preferences.get("avoid_defense_companies", False),
        not bool(constraints.get("is_defense_military_clearance")),
        str(constraints.get("constraint_matched_phrases", {}).get("defense_military_clearance", [])),
    )
    company_size_preference = advanced_preferences.get("company_size_preference", "No preference")
    add_check(
        "Company size preference",
        company_size_preference in {"Exclude tiny startups", "Large companies only"},
        bool(constraints.get("passes_company_size_filter", True)),
        str(constraints.get("company_size_match_text", "")),
    )
    return checks


def evaluate_constraint_passes(
    constraints: dict[str, object],
    advanced_preferences: dict,
) -> dict[str, object]:
    max_years = advanced_preferences.get("max_required_years", "No limit")
    required_years = constraints.get("required_years_min")
    max_allowed = safe_numeric_or_none(advanced_preferences.get("max_required_years_value"))
    if max_allowed is None:
        max_allowed = safe_numeric_or_none(str(max_years).split()[0])
    parsed_required_years = safe_numeric_or_none(required_years)
    if max_years == "No limit" or max_allowed is None:
        passes_years_filter = True
    else:
        passes_years_filter = (
            parsed_required_years is None or parsed_required_years <= max_allowed
        )

    passes_contract_filter = not (
        avoid_contract_temp_active(advanced_preferences)
        and bool(constraints.get("is_contract_or_temp"))
    )
    passes_unpaid_filter = not (
        avoid_unpaid_commission_active(advanced_preferences)
        and bool(constraints.get("is_unpaid_or_commission"))
    )
    seniority_preference = advanced_preferences.get("seniority_preference", "No preference")
    is_senior = bool(constraints.get("is_senior_staff_principal"))
    is_junior = bool(constraints.get("is_junior_entry_level"))
    is_manager = bool(constraints.get("is_manager_director"))
    is_internship = str(constraints.get("employment_type_inferred", "")).lower() == "internship"
    if seniority_preference == "Entry-level / New Grad only":
        passes_seniority_filter = not (is_senior or is_manager)
    elif seniority_preference == "Junior to Mid-level":
        passes_seniority_filter = not (is_senior or is_manager)
    elif seniority_preference == "Exclude Senior/Staff/Principal":
        passes_seniority_filter = not (is_senior or is_manager)
    elif seniority_preference in {"Mid to Senior allowed", "Senior+ preferred", "Exclude Junior/Entry-level"}:
        passes_seniority_filter = not is_junior
    else:
        passes_seniority_filter = True
    passes_manager_filter = not (
        seniority_preference in {
            "Entry-level / New Grad only",
            "Junior to Mid-level",
            "Exclude Senior/Staff/Principal",
        }
        and is_manager
    )
    passes_internship_filter = not (
        avoid_internship_active(advanced_preferences)
        and is_internship
    )
    minimum_salary = safe_numeric_or_none(advanced_preferences.get("minimum_salary"))
    salary_strictness = advanced_preferences.get("salary_strictness", "Soft")
    salary_number = safe_numeric_or_none(constraints.get("salary_number"))
    if not minimum_salary or minimum_salary <= 0 or salary_strictness != "Strict":
        passes_salary_filter = True
    else:
        salary_unknown = salary_number is None or salary_number <= 0
        passes_salary_filter = salary_unknown or salary_number >= minimum_salary
    passes_defense_filter = not (
        advanced_preferences.get("avoid_defense_companies", False)
        and bool(constraints.get("is_defense_military_clearance"))
    )
    company_size_preference = advanced_preferences.get("company_size_preference", "No preference")
    is_tiny = bool(constraints.get("is_tiny_startup"))
    is_small = bool(constraints.get("is_small_company"))
    if company_size_preference == "Exclude tiny startups":
        passes_company_size_filter = not is_tiny
    elif company_size_preference == "Large companies only":
        passes_company_size_filter = not (is_tiny or is_small)
    else:
        passes_company_size_filter = True
    passes_location_filter = True
    if advanced_preferences.get("location_strictness", "Balanced") == "Strict":
        passes_location_filter = bool(constraints.get("passes_location_filter", True))
    role_fit_score = safe_numeric_or_none(constraints.get("role_fit_score"))
    passes_role_fit_filter = not (
        advanced_preferences.get("role_fit_strictness", "Balanced") == "Strict"
        and role_fit_score is not None
        and role_fit_score < 0
    )

    active_filter_names: list[str] = []
    if max_years != "No limit":
        active_filter_names.append("max_required_years")
    if minimum_salary and minimum_salary > 0:
        active_filter_names.append("minimum_salary")
    if avoid_contract_temp_active(advanced_preferences):
        active_filter_names.append("avoid_contract_temp")
    if avoid_unpaid_commission_active(advanced_preferences):
        active_filter_names.append("avoid_unpaid_commission")
    if seniority_preference != "No preference":
        active_filter_names.append("seniority_preference")
    if avoid_internship_active(advanced_preferences):
        active_filter_names.append("avoid_internship")
    if advanced_preferences.get("avoid_defense_companies", False):
        active_filter_names.append("avoid_defense_military")
    if company_size_preference in {"Exclude tiny startups", "Large companies only"}:
        active_filter_names.append("company_size_preference")
    if advanced_preferences.get("location_strictness", "Balanced") == "Strict":
        active_filter_names.append("strict_location")
    if advanced_preferences.get("role_fit_strictness", "Balanced") == "Strict":
        active_filter_names.append("strict_role_fit")

    passes_all = all(
        [
            passes_years_filter,
            passes_contract_filter,
            passes_unpaid_filter,
            passes_seniority_filter,
            passes_manager_filter,
            passes_internship_filter,
            passes_salary_filter,
            passes_defense_filter,
            passes_company_size_filter,
            passes_location_filter,
            passes_role_fit_filter,
        ]
    )

    return {
        "passes_years_filter": passes_years_filter,
        "passes_contract_filter": passes_contract_filter,
        "passes_unpaid_filter": passes_unpaid_filter,
        "passes_seniority_filter": passes_seniority_filter,
        "passes_manager_filter": passes_manager_filter,
        "passes_internship_filter": passes_internship_filter,
        "passes_salary_filter": passes_salary_filter,
        "passes_defense_filter": passes_defense_filter,
        "passes_company_size_filter": passes_company_size_filter,
        "passes_role_fit_filter": passes_role_fit_filter,
        "passes_all_active_constraints": passes_all,
        "active_constraint_filters": ", ".join(active_filter_names) or "None",
        "active_constraint_penalty": 0 if passes_all else 500,
    }


def detect_role_family(target_role: str) -> str:
    normalized_target = normalize_search_text(target_role)
    if find_terms(normalized_target, ROLE_FAMILIES["mlops_infra"]["target_terms"]):
        return "mlops_infra"
    for family, config in ROLE_FAMILIES.items():
        if find_terms(normalized_target, config["target_terms"]):
            return family
    return "general"


def score_role_fit(target_role: str, title: str, description: str) -> tuple[int, str]:
    role_family = detect_role_family(target_role)
    if role_family == "general":
        return 0, "Role fit: +0 no specialized role family detected"

    config = ROLE_FAMILIES[role_family]
    title_text = str(title)
    searchable_text = f"{title} {description}"
    title_matches = find_terms(title_text, config["positive_terms"])
    positive_matches = find_terms(searchable_text, config["positive_terms"])
    negative_matches = find_terms(searchable_text, config["negative_terms"])
    non_job_title_matches = find_terms(title_text, NON_JOB_TRAINING_ROLE_TERMS)
    generic_software_title_matches = find_terms(title_text, GENERIC_SOFTWARE_TITLE_TERMS)

    score = 0
    reasons: list[str] = []

    if title_matches:
        title_bonus = min(50, 25 + 5 * len(title_matches))
        score += title_bonus
        reasons.append(
            f"Role fit: +{title_bonus} {role_family} title terms: {', '.join(title_matches[:5])}"
        )

    description_only_matches = [
        term for term in positive_matches if term not in title_matches
    ]
    if description_only_matches:
        description_bonus = min(35, 7 * len(description_only_matches))
        score += description_bonus
        reasons.append(
            f"Role fit: +{description_bonus} {role_family} content terms: "
            f"{', '.join(description_only_matches[:6])}"
        )

    if negative_matches:
        negative_penalty = min(60, 30 + 10 * (len(negative_matches) - 1))
        score -= negative_penalty
        reasons.append(
            f"Role fit: -{negative_penalty} weak/unrelated terms: {', '.join(negative_matches[:6])}"
        )

    if non_job_title_matches and role_family in {"ml_ai", "mlops_infra", "research"}:
        score -= 120
        reasons.append(
            "Role fit: -120 course/training/evaluator title: "
            f"{', '.join(non_job_title_matches[:5])}"
        )

    if (
        generic_software_title_matches
        and role_family in {"ml_ai", "mlops_infra", "research"}
        and not title_matches
    ):
        score -= 70
        reasons.append(
            "Role fit: -70 generic software title without target-role evidence: "
            f"{', '.join(generic_software_title_matches[:5])}"
        )

    if not positive_matches:
        score -= 35
        reasons.append(f"Role fit: -35 no clear {role_family} evidence")

    return score, "; ".join(reasons)


def score_role_fit_strictness(
    role_fit_score: int,
    strictness: str,
) -> tuple[int, str]:
    if strictness == "Flexible":
        return 0, ""
    if strictness == "Balanced" and role_fit_score < 10:
        return 20, "-20 role fit strictness penalty: weak match for target role"
    if strictness == "Strict":
        if role_fit_score < 0:
            return 60, "-60 role fit strictness penalty: poor match for target role"
        if role_fit_score < 15:
            return 40, "-40 role fit strictness penalty: weak match for target role"
    return 0, ""


def score_advanced_preferences(
    job: pd.Series,
    minimum_salary: int,
    advanced_preferences: dict,
) -> tuple[int, str]:
    title = str(job.get("title", ""))
    company = str(job.get("company", ""))
    description = str(job.get("description", ""))
    salary_number = safe_numeric_or_none(job.get("salary_number", 0)) or 0
    title_text = f"{title}"
    searchable_text = f"{title} {company} {description}"
    constraints = extract_job_constraints(job)
    matched_phrases = constraints["constraint_matched_phrases"]

    points = 0
    reasons: list[str] = []

    seniority_preference = advanced_preferences.get("seniority_preference", "No preference")
    seniority_points, seniority_reason = score_seniority_preference(
        constraints,
        seniority_preference,
        title,
    )
    points += seniority_points
    if seniority_reason:
        reasons.append(seniority_reason)

    max_years = advanced_preferences.get("max_required_years", "No limit")
    max_allowed_years = safe_numeric_or_none(advanced_preferences.get("max_required_years_value"))
    if max_allowed_years is None:
        max_allowed_years = safe_numeric_or_none(str(max_years).split()[0])
    if max_years != "No limit" and max_allowed_years is not None:
        required_years = safe_numeric_or_none(constraints["required_years_min"])
        if required_years is not None and required_years > max_allowed_years:
            points -= 45
            reasons.append(
                f"Advanced filter: -45 requires {required_years:g}+ years "
                f"over {max_allowed_years:g}-year maximum"
            )

    if avoid_contract_temp_active(advanced_preferences):
        matches = matched_phrases["contract_or_temp"]
        if constraints["is_contract_or_temp"]:
            points -= 80
            reasons.append(f"Advanced filter: -80 contract/temp terms: {', '.join(matches)}")

    if avoid_unpaid_commission_active(advanced_preferences):
        matches = matched_phrases["unpaid_or_commission"]
        if constraints["is_unpaid_or_commission"]:
            points -= 90
            reasons.append(
                f"Advanced filter: -90 unpaid/commission-only terms: {', '.join(matches)}"
            )

    if advanced_preferences.get("avoid_defense_companies"):
        matches = matched_phrases["defense_military_clearance"]
        if constraints["is_defense_military_clearance"]:
            points -= 55
            reasons.append(
                f"Advanced filter: -55 defense/military terms: {', '.join(matches)}"
            )

    company_size_preference = advanced_preferences.get("company_size_preference", "No preference")
    company_size = str(constraints.get("company_size_inferred", "Unknown"))
    company_size_match = str(constraints.get("company_size_match_text", ""))
    is_tiny = bool(constraints.get("is_tiny_startup"))
    is_small = bool(constraints.get("is_small_company"))
    is_medium = bool(constraints.get("is_medium_company"))
    is_large = bool(constraints.get("is_large_company"))
    is_enterprise = bool(constraints.get("is_enterprise_company"))
    if company_size_preference == "Prefer small companies":
        if is_tiny or is_small:
            points += 12
            reasons.append(f"Company size: +12 small-company signal: {company_size_match or company_size}")
        elif is_large or is_enterprise:
            points -= 8
            reasons.append("Company size: -8 larger-company signal")
    elif company_size_preference == "Prefer medium companies":
        if is_medium:
            points += 12
            reasons.append(f"Company size: +12 medium-company signal: {company_size_match or company_size}")
        elif is_tiny:
            points -= 10
            reasons.append("Company size: -10 tiny-startup signal")
    elif company_size_preference == "Prefer large companies":
        if is_large or is_enterprise:
            points += 18
            reasons.append(f"Company size: +18 large/enterprise signal: {company_size_match or company_size}")
        elif is_tiny or is_small:
            points -= 15
            reasons.append(f"Company size: -15 tiny/small signal: {company_size_match or company_size}")
    elif company_size_preference == "Exclude tiny startups" and is_tiny:
        points -= 100
        reasons.append(f"Company size: -100 excluded tiny startup: {company_size_match or company_size}")
    elif company_size_preference == "Large companies only":
        if is_large or is_enterprise:
            points += 15
            reasons.append(f"Company size: +15 large/enterprise signal: {company_size_match or company_size}")
        elif is_tiny or is_small:
            points -= 120
            reasons.append(f"Company size: -120 not large-company compatible: {company_size_match or company_size}")

    visa_preference = advanced_preferences.get("visa_sponsorship", "No preference")
    if visa_preference != "No preference":
        positive_matches = matched_phrases["visa_sponsorship"]
        negative_matches = matched_phrases["no_sponsorship"]
        contract_matches = matched_phrases["contract_or_temp"]
        if visa_preference == "Sponsorship required":
            if positive_matches:
                points += 10
                reasons.append(
                    f"Advanced filter: +10 sponsor-friendly terms: {', '.join(positive_matches[:4])}"
                )
            if negative_matches:
                points -= 50
                reasons.append(
                    f"Advanced filter: -50 sponsorship-negative terms: {', '.join(negative_matches[:3])}"
                )
            if contract_matches:
                points -= 20
                reasons.append(
                    f"Advanced filter: -20 contract/temp less sponsor-friendly: {', '.join(contract_matches[:3])}"
                )
        elif visa_preference == "Prefer sponsor-friendly" and positive_matches:
            points += 8
            reasons.append(
                f"Advanced filter: +8 sponsor-friendly terms: {', '.join(positive_matches[:4])}"
            )

    company_preference = advanced_preferences.get("company_preference", "No preference")
    company_terms = COMPANY_PREFERENCE_TERMS.get(company_preference)
    if company_terms:
        bonus_matches = find_terms(searchable_text, company_terms["bonus"])
        penalty_matches = find_terms(searchable_text, company_terms["penalty"])
        if bonus_matches:
            points += 12
            reasons.append(
                f"Advanced filter: +12 company preference match: {', '.join(bonus_matches[:5])}"
            )
        if penalty_matches:
            points -= 25
            reasons.append(
                f"Advanced filter: -25 company preference mismatch: {', '.join(penalty_matches[:5])}"
            )

    salary_strictness = advanced_preferences.get("salary_strictness", "Soft")
    if minimum_salary > 0:
        salary_known = salary_number > 0
        if salary_known and salary_number < minimum_salary:
            penalty_by_strictness = {"Soft": 10, "Medium": 25, "Strict": 80}
            penalty = penalty_by_strictness.get(salary_strictness, 10)
            points -= penalty
            reasons.append(
                f"Advanced filter: -{penalty} salary below minimum with {salary_strictness} strictness"
            )
        elif not salary_known and salary_strictness == "Medium":
            points -= 10
            reasons.append("Advanced filter: -10 salary unknown with Medium strictness")

    return points, "; ".join(reasons)


def get_data_path() -> Path:
    return CLEANED_DATA_PATH if CLEANED_DATA_PATH.exists() else SAMPLE_DATA_PATH


@st.cache_data
def load_jobs(path: Path) -> pd.DataFrame:
    jobs = pd.read_csv(path)
    for column in [
        "job_id",
        "title",
        "company",
        "location",
        "raw_location",
        "location_source",
        "employment_type",
        "description",
        "apply_link",
    ]:
        if column not in jobs.columns:
            jobs[column] = ""
        jobs[column] = jobs[column].fillna("").astype(str)
    jobs["raw_location"] = jobs["raw_location"].where(
        jobs["raw_location"].str.strip() != "",
        jobs["location"],
    )
    jobs["location_source"] = jobs.apply(
        lambda row: infer_location_source(row),
        axis=1,
    )
    if "salary" not in jobs.columns:
        jobs["salary"] = ""
    jobs["salary_text"] = jobs["salary"].apply(make_salary_text)
    jobs["salary_number"] = jobs["salary"].apply(parse_salary)
    return jobs


def score_job(
    job: pd.Series,
    target_role: str,
    skills: list[str],
    preferred_location: str,
    minimum_salary: int,
    dealbreakers: list[str],
    advanced_preferences: dict | None = None,
) -> tuple[int, str]:
    score = 0
    reasons: list[str] = []
    advanced_preferences = advanced_preferences or {}

    title = str(job["title"]).lower()
    description = str(job["description"]).lower()
    searchable_text = f"{title} {description}"

    role_points, role_reason = score_target_role_relevance(
        target_role=target_role,
        title=title,
        description=description,
    )
    score += role_points
    if role_reason:
        reasons.append(role_reason)

    matched_skills = [skill for skill in skills if skill in searchable_text]
    if matched_skills:
        points = min(40, 10 * len(matched_skills))
        score += points
        reasons.append(f"Matched skills: +{points} overlap: {', '.join(matched_skills)}")
    elif skills:
        reasons.append("Matched skills: +0 no listed skills found in posting")

    location_signals = extract_location_signals(job)
    location_points, location_reason = score_location_signals(
        preferred_location,
        location_signals,
        strictness=advanced_preferences.get("location_strictness", "Balanced"),
    )
    score += location_points
    if location_reason:
        reasons.append(location_reason)

    if minimum_salary > 0:
        if int(job["salary_number"]) >= minimum_salary:
            points = 15
            score += points
            reasons.append(f"Salary fit: +{points} appears to meet minimum")
        else:
            reasons.append("Salary fit: +0 unknown or below minimum")

    dealbreaker_points, dealbreaker_reason = score_dealbreakers(
        dealbreakers=dealbreakers,
        title=title,
        description=description,
    )
    score += dealbreaker_points
    if dealbreaker_reason:
        reasons.append(dealbreaker_reason)

    advanced_points, advanced_reason = score_advanced_preferences(
        job=job,
        minimum_salary=minimum_salary,
        advanced_preferences=advanced_preferences,
    )
    score += advanced_points
    if advanced_reason:
        reasons.append(advanced_reason)

    if not reasons:
        reasons.append("No preferences entered yet, so this job keeps a neutral score.")

    return score, "; ".join(reasons)


def rank_jobs(
    jobs: pd.DataFrame,
    target_role: str,
    skills_text: str,
    preferred_location: str,
    minimum_salary: int,
    dealbreakers_text: str,
    advanced_preferences: dict | None = None,
) -> pd.DataFrame:
    skills = split_keywords(skills_text)
    dealbreakers = split_keywords(dealbreakers_text)

    ranked_jobs = jobs.copy()
    scores_and_reasons = ranked_jobs.apply(
        lambda job: score_job(
            job,
            target_role=target_role,
            skills=skills,
            preferred_location=preferred_location,
            minimum_salary=minimum_salary,
            dealbreakers=dealbreakers,
            advanced_preferences=advanced_preferences,
        ),
        axis=1,
    )

    ranked_jobs["score"] = [item[0] for item in scores_and_reasons]
    ranked_jobs["explanation"] = [item[1] for item in scores_and_reasons]
    ranked_jobs["matched_skills"] = ranked_jobs.apply(
        lambda job: ", ".join(
            skill
            for skill in skills
            if skill in f"{str(job.get('title', '')).lower()} {str(job.get('description', '')).lower()}"
        ),
        axis=1,
    )
    role_fit_scores_and_reasons = ranked_jobs.apply(
        lambda job: score_role_fit(
            target_role=target_role,
            title=str(job.get("title", "")),
            description=str(job.get("description", "")),
        ),
        axis=1,
    )
    ranked_jobs["role_fit_score"] = [item[0] for item in role_fit_scores_and_reasons]
    ranked_jobs["role_fit_reason"] = [item[1] for item in role_fit_scores_and_reasons]
    constraint_rows = ranked_jobs.apply(extract_job_constraints, axis=1)
    constraint_columns = [
        "required_years_min",
        "years_requirement_text",
        "employment_type_inferred",
        "is_contract_or_temp",
        "is_unpaid_or_commission",
        "contract_temp_text",
        "unpaid_commission_text",
        "is_senior_staff_principal",
        "is_junior_entry_level",
        "is_manager_director",
        "is_defense_military_clearance",
        "defense_match_text",
        "company_size_inferred",
        "company_size_confidence",
        "company_size_match_text",
        "is_tiny_startup",
        "is_small_company",
        "is_medium_company",
        "is_large_company",
        "is_enterprise_company",
        "has_visa_sponsorship_signal",
        "has_no_sponsorship_signal",
        "raw_location",
        "location_source",
        "location_confidence",
        "is_us_location",
        "is_remote",
        "is_us_remote_location",
        "has_explicit_non_us_location_signal",
        "is_hybrid",
        "is_california",
        "is_bay_area",
        "is_california_or_bay_area",
        "is_non_us_location",
        "location_matched_phrases",
        "role_family_inferred",
        "constraint_matched_phrases",
    ]
    for column in constraint_columns:
        ranked_jobs[column] = [constraints[column] for constraints in constraint_rows]
    location_strictness = (advanced_preferences or {}).get(
        "location_strictness", "Balanced"
    )
    location_checks = ranked_jobs.apply(
        lambda job: evaluate_location_preference(
            preferred_location,
            job.to_dict(),
            strictness=location_strictness,
        ),
        axis=1,
    )
    ranked_jobs["passes_location_filter"] = [item[0] for item in location_checks]
    ranked_jobs["location_filter_reason"] = [item[1] for item in location_checks]
    pass_rows = ranked_jobs.apply(
        lambda job: evaluate_constraint_passes(
            job.to_dict(),
            advanced_preferences or {},
        ),
        axis=1,
    )
    pass_columns = [
        "passes_years_filter",
        "passes_contract_filter",
        "passes_unpaid_filter",
        "passes_seniority_filter",
        "passes_manager_filter",
        "passes_internship_filter",
        "passes_salary_filter",
        "passes_defense_filter",
        "passes_company_size_filter",
        "passes_role_fit_filter",
        "passes_all_active_constraints",
        "active_constraint_filters",
        "active_constraint_penalty",
    ]
    for column in pass_columns:
        ranked_jobs[column] = [passes[column] for passes in pass_rows]
    ranked_jobs["seniority_preference"] = (advanced_preferences or {}).get(
        "seniority_preference",
        "No preference",
    )
    ranked_jobs["preferred_location"] = preferred_location
    ranked_jobs["location_strictness"] = (advanced_preferences or {}).get(
        "location_strictness",
        "Balanced",
    )
    ranked_jobs["company_size_preference"] = (advanced_preferences or {}).get(
        "company_size_preference",
        "No preference",
    )
    ranked_jobs["salary_minimum"] = minimum_salary
    ranked_jobs["salary_strictness"] = (advanced_preferences or {}).get(
        "salary_strictness",
        "Soft",
    )
    ranked_jobs["max_required_years_setting"] = (advanced_preferences or {}).get(
        "max_required_years",
        "No limit",
    )
    ranked_jobs["required_years_min"] = ranked_jobs["required_years_min"].apply(
        lambda value: "Unknown" if pd.isna(value) else int(value)
    )
    role_fit_strictness = (advanced_preferences or {}).get(
        "role_fit_strictness", "Balanced"
    )
    strictness_scores_and_reasons = ranked_jobs["role_fit_score"].apply(
        lambda score: score_role_fit_strictness(int(score), role_fit_strictness)
    )
    ranked_jobs["role_fit_strictness_penalty"] = [
        item[0] for item in strictness_scores_and_reasons
    ]
    ranked_jobs["role_fit_strictness_reason"] = [
        item[1] for item in strictness_scores_and_reasons
    ]
    if "embedding_similarity" in ranked_jobs.columns:
        ranked_jobs["embedding_similarity"] = pd.to_numeric(
            ranked_jobs["embedding_similarity"], errors="coerce"
        ).fillna(0.0)
    else:
        ranked_jobs["embedding_similarity"] = 0.0

    ranked_jobs["embedding_bonus"] = (
        ranked_jobs["embedding_similarity"] * EMBEDDING_BONUS_SCALE
    ).round(2)
    ranked_jobs["final_score"] = (
        ranked_jobs["score"]
        + ranked_jobs["role_fit_score"]
        + ranked_jobs["embedding_bonus"]
        - ranked_jobs["role_fit_strictness_penalty"]
        - ranked_jobs["active_constraint_penalty"]
    ).round(2)
    ranked_jobs["feedback_label"] = "Skip"
    ranked_jobs["feedback_adjustment_score"] = 0
    ranked_jobs["feedback_adjustment_reason"] = ""
    ranked_jobs["adjusted_final_score"] = ranked_jobs["final_score"]
    ranked_jobs["explanation"] = ranked_jobs.apply(
        lambda job: (
            f"Embedding similarity contribution: +{job['embedding_bonus']:.2f} "
            f"(similarity {job['embedding_similarity']:.3f}). "
            f"{job['role_fit_reason']}. "
            + (
                f"{job['role_fit_strictness_reason']}. "
                if job["role_fit_strictness_reason"]
                else ""
            )
            + (
                "Active constraint penalty: -500 failed one or more active filters. "
                if job["active_constraint_penalty"]
                else ""
            )
            + job["explanation"]
        ),
        axis=1,
    )

    sort_columns = [
        "passes_all_active_constraints",
        "final_score",
        "role_fit_score",
        "score",
        "embedding_similarity",
        "salary_number",
        "title",
    ]
    sort_ascending = [False, False, False, False, False, False, True]
    sorted_jobs = ranked_jobs.sort_values(
        by=sort_columns,
        ascending=sort_ascending,
    ).reset_index(drop=True)
    sorted_jobs.insert(0, "rank", range(1, len(sorted_jobs) + 1))
    return sorted_jobs


def get_embedding_candidates(
    jobs: pd.DataFrame,
    target_role: str,
    combined_skills_text: str,
    resume_text: str,
) -> tuple[pd.DataFrame, bool, str]:
    if not embedding_files_exist():
        return jobs.copy(), False, "Embedding files missing; used full rule-based ranking."

    profile_text = build_user_profile_text(
        target_role=target_role,
        combined_skills_text=combined_skills_text,
        resume_text=resume_text,
    )
    candidates = retrieve_top_candidates(
        profile_text,
        top_k=EMBEDDING_CANDIDATE_COUNT,
    )
    retrieval_method = candidates.attrs.get("retrieval_method", "bruteforce")
    candidates = candidates[candidates["row_index"].between(0, len(jobs) - 1)]

    if candidates.empty:
        return jobs.copy(), False, "No embedding candidates found; used full rule-based ranking."

    candidate_jobs = jobs.iloc[candidates["row_index"].tolist()].copy()
    candidate_jobs["embedding_similarity"] = candidates["embedding_similarity"].to_list()
    search_label = (
        "ANN vector search over job embeddings"
        if retrieval_method == "ann"
        else "brute-force embedding similarity"
    )
    note = (
        f"Retrieved {len(candidate_jobs):,} candidate jobs using {search_label}, "
        "then applied rule-based ranking."
    )
    return candidate_jobs, True, note


def prepare_user_download_jobs(jobs: pd.DataFrame) -> pd.DataFrame:
    download_jobs = jobs.copy()

    if "apply_url" not in download_jobs.columns:
        for link_column in ["apply_link", "url", "link"]:
            if link_column in download_jobs.columns:
                download_jobs["apply_url"] = download_jobs[link_column]
                break
        else:
            download_jobs["apply_url"] = ""

    if "match_explanation" not in download_jobs.columns:
        def make_match_explanation(row: pd.Series) -> str:
            parts: list[str] = []
            matched_skills = str(row.get("matched_skills", "")).strip()
            role_family = str(row.get("role_family_inferred", "")).strip()
            constraints_passed = row.get("passes_all_active_constraints", "")
            if matched_skills:
                parts.append(f"Matched skills: {matched_skills}")
            if role_family and role_family.lower() != "unknown":
                parts.append(f"Role family: {role_family}")
            if constraints_passed != "":
                status = "passed" if bool(constraints_passed) else "did not pass"
                parts.append(f"Active constraints: {status}")
            return "; ".join(parts) or str(row.get("explanation", "Ranked by JobPilot."))

        download_jobs["match_explanation"] = download_jobs.apply(
            make_match_explanation,
            axis=1,
        )

    return download_jobs


def clean_candidate_name(candidate_name: str) -> str:
    """Clean persona-card headers without altering ordinary candidate names."""
    original_name = re.sub(r"\s+", " ", str(candidate_name or "")).strip()
    if not original_name or original_name.lower() in {"unknown", "nan", "none"}:
        return "[Add your name]"

    if re.match(r"^persona(?:\s+\d+)?\b", original_name, flags=re.IGNORECASE):
        parts = [
            part.strip()
            for part in re.split(r"\s+(?:—|–|-)\s+", original_name)
            if part.strip()
        ]
        if len(parts) >= 2 and parts[1]:
            return parts[1]

    return original_name


def extract_candidate_background(resume_text: str) -> str:
    """Extract a concise, evidence-based background summary from resume text."""
    fallback = "[Add a 1-2 sentence background summary from your resume.]"
    if not resume_text:
        return fallback

    excluded_headings = {
        "background",
        "candidate name",
        "dealbreakers",
        "pass criteria",
        "persona",
        "preferences",
        "skills",
        "target role",
        "target roles",
    }
    section_headings = excluded_headings | {
        "certifications",
        "education",
        "experience",
        "projects",
        "summary",
        "work experience",
    }
    evidence_patterns = [
        r"\bbackground\b",
        r"\b(?:bachelor|bachelors|bs|ba|master|masters|ms|msba|mba|phd|degree)\b",
        r"\b(?:university|college|school|graduate|graduating|student)\b",
        r"\b\d+\+?\s+years?\b",
        r"\b(?:year|years)\s+(?:of|in)\s+experience\b",
        r"\b(?:intern|internship|internships)\b",
        r"\b(?:career transition|career pivot|pivot|transition|move to|moving to)\b",
        r"\b(?:opt|h-?1b|visa|sponsorship)\b",
        r"\b(?:engineer|analyst|scientist|researcher|developer|manager)\b",
    ]

    lines = [
        re.sub(r"\s+", " ", line).strip(" -•\t")
        for line in resume_text.replace("\r", "\n").splitlines()
        if re.sub(r"\s+", " ", line).strip()
    ]

    background_lines: list[str] = []
    in_background_section = False
    for line in lines:
        normalized = line.rstrip(":").strip().lower()
        if re.match(r"^persona(?:\s+\d+)?\b", line, flags=re.IGNORECASE):
            continue
        if normalized in section_headings:
            in_background_section = normalized in {"background", "summary"}
            continue
        if any(
            normalized.startswith(f"{heading}:")
            for heading in excluded_headings - {"background"}
        ):
            in_background_section = False
            continue
        if normalized.startswith("background:"):
            content = line.split(":", 1)[1].strip()
            if content:
                background_lines.append(content)
            in_background_section = True
            continue
        if in_background_section:
            background_lines.append(line)
            if len(background_lines) >= 3:
                break

    if not background_lines:
        for line in lines[1:20]:
            normalized = line.rstrip(":").strip().lower()
            if normalized in excluded_headings or re.match(
                r"^persona(?:\s+\d+)?\b",
                line,
                flags=re.IGNORECASE,
            ):
                continue
            if any(re.search(pattern, line, flags=re.IGNORECASE) for pattern in evidence_patterns):
                background_lines.append(line)
            if len(background_lines) >= 3:
                break

    cleaned_lines: list[str] = []
    seen_lines: set[str] = set()
    for line in background_lines:
        cleaned = re.sub(r"^(?:background|summary)\s*:\s*", "", line, flags=re.IGNORECASE)
        cleaned = cleaned.strip()
        if not cleaned or cleaned.lower() in excluded_headings:
            continue
        if cleaned.lower() not in seen_lines:
            cleaned_lines.append(cleaned)
            seen_lines.add(cleaned.lower())

    return " ".join(cleaned_lines)[:500].strip() or fallback


def is_noisy_resume_evidence_line(
    line: str,
    candidate_name: str,
    background: str,
    known_skills: list[str],
) -> bool:
    """Return True for resume structure/content that is not project evidence."""
    normalized = re.sub(r"\s+", " ", line).strip(" -•\t")
    normalized_lower = normalized.rstrip(":").lower()
    excluded_labels = {
        "background",
        "candidate name",
        "certifications",
        "dealbreakers",
        "education",
        "experience",
        "pass criteria",
        "preferences",
        "projects",
        "summary",
        "skills",
        "target role",
        "target roles",
        "work experience",
    }
    if not normalized:
        return True
    if re.match(r"^persona(?:\s+\d+)?\b", normalized, flags=re.IGNORECASE):
        return True
    if normalized_lower in excluded_labels:
        return True
    if any(normalized_lower.startswith(f"{label}:") for label in excluded_labels):
        return True
    if normalized_lower == candidate_name.strip().lower():
        return True
    if background and normalized_lower in background.lower():
        return True

    comma_parts = [part.strip().lower() for part in normalized.split(",") if part.strip()]
    known_skill_keys = {skill.lower() for skill in known_skills}
    if len(comma_parts) >= 2 and sum(part in known_skill_keys for part in comma_parts) >= 2:
        return True
    return False


def build_project_suggestions(skills: list[str]) -> list[str]:
    """Create cautious, skill-based project prompts without inventing experience."""
    prioritized = skills[:6]
    primary = prioritized[:4]
    primary_text = ", ".join(primary[:-1]) + f", or {primary[-1]}" if len(primary) > 1 else (
        primary[0] if primary else "the selected job's core skills"
    )
    lower_skills = {skill.lower() for skill in prioritized}
    suggestions = [
        (
            "- Highlight an internship, academic project, or resume-backed example "
            f"using {primary_text} to solve a relevant problem."
        ),
    ]

    dashboard_skills = [
        skill
        for skill in prioritized
        if skill.lower() in {"tableau", "power bi", "excel", "dashboarding", "business intelligence"}
    ]
    if dashboard_skills:
        suggestions.append(
            "- Add a concrete example of building dashboards, reports, or BI "
            f"deliverables using {', '.join(dashboard_skills)} or related tools."
        )

    if lower_skills & {
        "sql",
        "spark",
        "pyspark",
        "etl",
        "data engineering",
        "pandas",
        "python",
        "r",
    }:
        suggestions.append(
            "- Describe a project where you cleaned, transformed, analyzed, or "
            "queried data to support a decision or technical outcome."
        )

    if len(suggestions) < 3:
        suggestions.append(
            f"- Add a concrete example from your resume showing {primary_text}."
        )
    suggestions.append(
        "- If accurate, quantify the project scope, such as dataset size, reporting "
        "frequency, deliverables created, or stakeholders supported."
    )
    return suggestions[:4]


def extract_candidate_skill_evidence_text(resume_text: str) -> str:
    """Remove target/preference sections before scanning resume text for skills."""
    if not resume_text:
        return ""

    excluded_sections = {
        "dealbreakers",
        "pass criteria",
        "preferences",
        "preference",
        "target role",
        "target roles",
        "tailoring notes",
    }
    known_sections = excluded_sections | {
        "background",
        "certifications",
        "education",
        "experience",
        "projects",
        "skills",
        "summary",
        "technical skills",
        "work experience",
    }
    evidence_lines: list[str] = []
    active_section = ""

    for raw_line in resume_text.replace("\r", "\n").splitlines():
        line = re.sub(r"\s+", " ", raw_line).strip()
        if not line or re.match(r"^persona(?:\s+\d+)?\b", line, flags=re.IGNORECASE):
            continue

        if ":" in line:
            possible_heading, content = line.split(":", 1)
            heading = possible_heading.strip().lower()
            if heading in known_sections:
                active_section = heading
                if heading not in excluded_sections and content.strip():
                    evidence_lines.append(content.strip())
                continue

        heading = line.rstrip(":").strip().lower()
        if heading in known_sections:
            active_section = heading
            continue
        if active_section in excluded_sections:
            continue
        evidence_lines.append(line)

    return "\n".join(evidence_lines)


def get_candidate_supported_skills(
    combined_skills_text: str,
    resume_text: str,
) -> set[str]:
    """Return normalized skills explicitly supported by candidate-provided evidence."""
    supported = {skill.lower() for skill in split_keywords(combined_skills_text)}
    known_skill_terms = {
        "airflow",
        "aws",
        "azure",
        "cloud",
        "cloud infrastructure",
        "data pipeline",
        "databricks",
        "docker",
        "feature pipeline",
        "gcp",
        "java",
        "kafka",
        "kubernetes",
        "microservices",
        "ml infrastructure",
        "ml platform",
        "mlflow",
        "mlops",
        "model deployment",
        "model serving",
        "production ml",
        "spark",
        "tensorflow serving",
        "terraform",
    }
    candidate_evidence_text = extract_candidate_skill_evidence_text(resume_text)
    normalized_resume = normalize_search_text(candidate_evidence_text)
    for term in known_skill_terms:
        if contains_phrase(normalized_resume, term):
            supported.add(term)
    return supported


def detect_ml_infrastructure_skills(
    skills: list[str],
    selected_job: pd.Series,
    target_role: str,
    resume_text: str,
) -> list[str]:
    """Detect infrastructure skills when the candidate/job has ML platform context."""
    infrastructure_terms = {
        "aws",
        "azure",
        "cloud",
        "cloud infrastructure",
        "airflow",
        "data pipeline",
        "databricks",
        "docker",
        "feature pipeline",
        "gcp",
        "kafka",
        "kubernetes",
        "microservices",
        "ml infrastructure",
        "ml platform",
        "mlflow",
        "mlops",
        "model deployment",
        "model serving",
        "production ml",
        "spark",
        "tensorflow serving",
        "terraform",
    }
    ml_context_terms = {
        "ai infrastructure",
        "machine learning infrastructure",
        "machine learning platform",
        "ml infrastructure",
        "ml platform",
        "mlops",
        "model deployment",
        "model serving",
        "production machine learning",
        "production ml",
    }
    context_text = " ".join(
        [
            target_role,
            str(selected_job.get("title", "")),
            str(selected_job.get("description", "")),
            str(selected_job.get("matched_skills", "")),
            str(selected_job.get("role_family_inferred", "")),
            resume_text,
        ]
    ).lower()
    has_ml_infrastructure_context = (
        str(selected_job.get("role_family_inferred", "")).lower() == "mlops_infra"
        or any(term in context_text for term in ml_context_terms)
    )
    if not has_ml_infrastructure_context:
        return []

    return [
        skill
        for skill in skills
        if skill.lower() in infrastructure_terms
    ]


def build_ml_infrastructure_suggestions(skills: list[str]) -> list[str]:
    """Create cautious project prompts framed around ML platform infrastructure."""
    skill_map = {skill.lower(): skill for skill in skills}
    streaming_skills = [
        skill_map[key]
        for key in ["kafka"]
        if key in skill_map
    ]
    processing_skills = [
        skill_map[key]
        for key in ["spark", "databricks", "airflow"]
        if key in skill_map
    ]
    platform_skills = [
        skill_map[key]
        for key in [
            "kubernetes",
            "docker",
            "microservices",
            "aws",
            "gcp",
            "azure",
            "cloud",
            "cloud infrastructure",
            "terraform",
            "mlflow",
            "tensorflow serving",
            "mlops",
            "ml platform",
            "ml infrastructure",
            "model deployment",
            "model serving",
            "production ml",
        ]
        if key in skill_map
    ]
    suggestions: list[str] = []
    if streaming_skills:
        suggestions.append(
            "- Highlight a resume-backed example where "
            f"{', '.join(streaming_skills)} was used for event streaming or scalable "
            "infrastructure, if accurate."
        )
    if processing_skills:
        suggestions.append(
            "- Describe a resume-backed project involving "
            f"{', '.join(processing_skills)} for large-scale data processing, if "
            "supported by your resume."
        )
    if platform_skills:
        suggestions.append(
            "- Add a concrete example from your resume showing "
            f"{', '.join(platform_skills)} in deploying or supporting reliable "
            "platform services."
        )
    if len(suggestions) < 3:
        suggestions.append(
            "- Connect infrastructure experience to the selected role only if that "
            "connection is supported by your resume."
        )
    suggestions.append(
        "- If accurate, quantify scale, reliability, latency, deployment frequency, "
        "or data volume using only details present in the original resume."
    )
    return suggestions[:4]


def generate_tailored_resume(
    candidate_name: str,
    resume_text: str,
    selected_job: pd.Series,
    matched_skills: str,
    target_role: str,
    combined_skills_text: str,
) -> str:
    """Create a cautious, template-based resume draft for one selected job."""
    job_title = str(selected_job.get("title", target_role or "Target Role")).strip()
    company = str(selected_job.get("company", "Target Company")).strip()
    description = str(selected_job.get("description", ""))
    role_family = str(selected_job.get("role_family_inferred", "Unknown")).strip()
    required_years = selected_job.get("required_years_min", "Unknown")
    salary_text = str(selected_job.get("salary_text", "Unknown")).strip()

    matched_keys = set(split_keywords(matched_skills))
    combined = split_keywords(combined_skills_text)
    candidate_supported_skills = get_candidate_supported_skills(
        combined_skills_text,
        resume_text,
    )
    relevant_skills: list[str] = []
    seen_skills: set[str] = set()
    for skill in combined + sorted(candidate_supported_skills):
        skill_key = skill.lower()
        if skill_key in candidate_supported_skills and skill_key not in seen_skills:
            relevant_skills.append(format_skill_name(skill))
            seen_skills.add(skill_key)

    job_keywords = [
        skill
        for skill in relevant_skills
        if skill.lower() in description.lower() or skill.lower() in matched_keys
    ][:10]
    if not job_keywords:
        job_keywords = relevant_skills[:10]
    ml_infrastructure_skills = detect_ml_infrastructure_skills(
        relevant_skills,
        selected_job,
        target_role,
        resume_text,
    )
    ml_infrastructure_keys = {skill.lower() for skill in ml_infrastructure_skills}
    general_skills = [
        skill for skill in relevant_skills if skill.lower() not in ml_infrastructure_keys
    ]

    cleaned_name = clean_candidate_name(candidate_name)
    name_lines = ["Candidate Name", cleaned_name]
    background = extract_candidate_background(resume_text)
    candidate_evidence_text = extract_candidate_skill_evidence_text(resume_text)
    resume_lines = [
        re.sub(r"\s+", " ", line).strip(" -•\t")
        for line in candidate_evidence_text.splitlines()
        if re.sub(r"\s+", " ", line).strip()
    ]
    evidence_action_pattern = re.compile(
        r"\b(?:analyzed|built|created|designed|developed|implemented|improved|"
        r"led|managed|modeled|optimized|project|published|researched|supported|"
        r"transformed|worked)\b",
        flags=re.IGNORECASE,
    )
    evidence_lines = [
        line
        for line in resume_lines
        if not is_noisy_resume_evidence_line(
            line,
            cleaned_name,
            background,
            relevant_skills,
        )
        and evidence_action_pattern.search(line)
        and any(skill.lower() in line.lower() for skill in job_keywords)
    ][:4]

    skill_lines = "\n".join(
        f"- {skill}{' (matched to job)' if skill.lower() in matched_keys else ''}"
        for skill in general_skills[:15]
    ) or "- Add verified skills from the original resume that match this job."
    ml_infrastructure_skill_lines = "\n".join(
        f"- {skill}{' (matched to job)' if skill.lower() in matched_keys else ''}"
        for skill in ml_infrastructure_skills[:15]
    )

    experience_suggestions = [
        f"- Refine and quantify this existing resume evidence where accurate: {line}"
        for line in evidence_lines[:2]
    ]
    fallback_suggestions = (
        build_ml_infrastructure_suggestions(ml_infrastructure_skills)
        if ml_infrastructure_skills
        else build_project_suggestions(job_keywords or relevant_skills)
    )
    quantify_suggestion = next(
        (
            suggestion
            for suggestion in fallback_suggestions
            if suggestion.startswith("- If accurate, quantify")
        ),
        "",
    )
    for suggestion in fallback_suggestions:
        if suggestion == quantify_suggestion:
            continue
        if suggestion not in experience_suggestions and len(experience_suggestions) < 3:
            experience_suggestions.append(suggestion)
    if quantify_suggestion:
        experience_suggestions.append(quantify_suggestion)
    experience_lines = "\n".join(experience_suggestions[:4])

    tailoring_notes = [
        f"- Emphasized target title: {job_title}",
        f"- Target company: {company}",
        f"- Inferred role family: {role_family or 'Unknown'}",
        f"- Required years detected: {required_years}",
        f"- Salary listed: {salary_text or 'Unknown'}",
    ]
    if job_keywords:
        tailoring_notes.append(f"- Prioritized job-relevant skills: {', '.join(job_keywords)}")
    tailoring_notes.append(
        "- Verify every statement against the original resume before submitting."
    )

    resume_sections = name_lines + [
            "",
            "Background",
            background,
            "",
            "Target Role",
            job_title,
            "",
            "Target Company",
            company,
            "",
            "Relevant Skills",
            skill_lines,
    ]
    if ml_infrastructure_skill_lines:
        resume_sections.extend(
            [
                "",
                "ML Infrastructure Skills",
                ml_infrastructure_skill_lines,
            ]
        )
    resume_sections.extend(
        [
            "",
            "Experience / Project Bullet Suggestions",
            experience_lines,
            "",
            "Tailoring Notes",
            "These notes are for review and can be deleted before submitting the resume.",
            "\n".join(tailoring_notes),
        ]
    )
    return "\n".join(resume_sections)


def format_skill_name(skill: str) -> str:
    skill_names = {
        "a/b testing": "A/B Testing",
        "airflow": "Airflow",
        "aws": "AWS",
        "gcp": "GCP",
        "etl": "ETL",
        "ml": "ML",
        "ml infrastructure": "ML Infrastructure",
        "ml platform": "ML Platform",
        "mlops": "MLOps",
        "model deployment": "Model Deployment",
        "model serving": "Model Serving",
        "nlp": "NLP",
        "numpy": "NumPy",
        "pandas": "pandas",
        "power bi": "Power BI",
        "pyspark": "PySpark",
        "pytorch": "PyTorch",
        "production ml": "Production ML",
        "scikit-learn": "scikit-learn",
        "sql": "SQL",
        "tensorflow": "TensorFlow",
        "tensorflow serving": "TensorFlow Serving",
        "terraform": "Terraform",
    }
    normalized = re.sub(r"\s+", " ", str(skill)).strip().lower()
    return skill_names.get(normalized, normalized.title())


def build_resume_docx(resume_draft: str) -> bytes:
    """Convert the editable plain-text resume draft into a DOCX byte stream."""
    from docx import Document

    section_headings = {
        "Candidate Name",
        "Background",
        "Target Role",
        "Target Company",
        "Relevant Skills",
        "ML Infrastructure Skills",
        "Experience / Project Bullet Suggestions",
        "Tailoring Notes",
    }
    document = Document()
    lines = resume_draft.splitlines()

    for index, raw_line in enumerate(lines):
        line = raw_line.strip()
        if not line:
            continue
        if index == 0 and line.startswith("Candidate Name:"):
            document.add_heading(line, level=1)
        elif line in section_headings:
            document.add_heading(line, level=1 if line == "Candidate Name" else 2)
        elif line.startswith("- "):
            document.add_paragraph(line[2:], style="List Bullet")
        else:
            document.add_paragraph(line)

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def stable_widget_key(*parts: object) -> str:
    raw_key = "_".join(str(part) for part in parts)
    return re.sub(r"[^a-zA-Z0-9_-]+", "_", raw_key).strip("_")[:180]


def split_comma_values(value: object) -> list[str]:
    if value is None or pd.isna(value):
        return []
    return [
        part.strip().lower()
        for part in str(value).split(",")
        if part.strip() and part.strip().lower() not in FEEDBACK_IGNORED_TERMS
    ]


def extract_title_terms(title: object) -> list[str]:
    words = re.findall(r"[a-zA-Z][a-zA-Z+#.]{2,}", str(title).lower())
    stop_words = {
        "and",
        "the",
        "for",
        "with",
        "remote",
        "senior",
        "junior",
        "lead",
        "staff",
        "principal",
    }
    return [
        word
        for word in words
        if word not in stop_words
        and word not in GENERIC_ROLE_WORDS
        and word not in FEEDBACK_IGNORED_TERMS
    ][:8]


def build_feedback_record(
    job: pd.Series,
    feedback_round: int,
    feedback_label: str,
) -> dict[str, object]:
    reward_map = {"Accept": 1, "Skip": 0, "Reject": -1}
    return {
        "feedback_round": feedback_round,
        "title": job.get("title", ""),
        "company": job.get("company", ""),
        "location": job.get("location", ""),
        "role_family_inferred": job.get("role_family_inferred", ""),
        "company_size_inferred": job.get("company_size_inferred", ""),
        "matched_skills": job.get("matched_skills", ""),
        "employment_type_inferred": job.get("employment_type_inferred", ""),
        "required_years_min": job.get("required_years_min", ""),
        "final_score": job.get("final_score", 0),
        "is_non_us_location": job.get("is_non_us_location", False),
        "passes_location_filter": job.get("passes_location_filter", True),
        "is_defense_military_clearance": job.get("is_defense_military_clearance", False),
        "is_contract_or_temp": job.get("is_contract_or_temp", False),
        "passes_all_active_constraints": job.get("passes_all_active_constraints", True),
        "feedback_label": feedback_label,
        "reward": reward_map.get(feedback_label, 0),
    }


def build_feedback_preferences(
    feedback_history: list[dict[str, object]],
    protected_role_family: str = "",
) -> dict[str, Counter]:
    preferences = {
        "accepted_role_family": Counter(),
        "accepted_skills": Counter(),
        "accepted_company_size": Counter(),
        "accepted_title_terms": Counter(),
        "rejected_role_family": Counter(),
        "rejected_skills": Counter(),
        "rejected_company_size": Counter(),
        "rejected_employment_type": Counter(),
        "rejected_hard_constraints": Counter(),
        "rejected_title_terms": Counter(),
    }
    protected_role_family = protected_role_family.strip().lower()

    for record in feedback_history:
        reward = int(record.get("reward", 0))
        if reward == 0:
            continue

        prefix = "accepted" if reward > 0 else "rejected"
        role_family = str(record.get("role_family_inferred", "")).strip().lower()
        company_size = str(record.get("company_size_inferred", "")).strip().lower()
        employment_type = str(record.get("employment_type_inferred", "")).strip().lower()
        if role_family and role_family != "unknown":
            if reward < 0 and role_family == protected_role_family:
                pass
            else:
                preferences[f"{prefix}_role_family"][role_family] += 1
        if company_size and company_size != "unknown":
            preferences[f"{prefix}_company_size"][company_size] += 1
        if (
            reward < 0
            and employment_type
            and employment_type != "unknown"
            and employment_type not in FEEDBACK_IGNORED_TERMS
        ):
            preferences["rejected_employment_type"][employment_type] += 1
        if reward < 0:
            if bool(record.get("is_non_us_location")) or not bool(record.get("passes_location_filter", True)):
                preferences["rejected_hard_constraints"]["non_us_location"] += 1
            if bool(record.get("is_defense_military_clearance")):
                preferences["rejected_hard_constraints"]["security_clearance_or_defense"] += 1
            if bool(record.get("is_contract_or_temp")):
                preferences["rejected_hard_constraints"]["contract_or_temp"] += 1
        for skill in split_comma_values(record.get("matched_skills", "")):
            preferences[f"{prefix}_skills"][skill] += 1
        for term in extract_title_terms(record.get("title", "")):
            preferences[f"{prefix}_title_terms"][term] += 1

    return preferences


def score_feedback_adjustment(
    job: pd.Series,
    feedback_preferences: dict[str, Counter],
) -> tuple[int, str]:
    score = 0
    reasons: list[str] = []
    role_family = str(job.get("role_family_inferred", "")).strip().lower()
    company_size = str(job.get("company_size_inferred", "")).strip().lower()
    employment_type = str(job.get("employment_type_inferred", "")).strip().lower()
    matched_skills = split_comma_values(job.get("matched_skills", ""))
    title_terms = extract_title_terms(job.get("title", ""))

    if role_family:
        accepted_count = feedback_preferences["accepted_role_family"][role_family]
        rejected_count = feedback_preferences["rejected_role_family"][role_family]
        if accepted_count:
            points = 10 * accepted_count
            score += points
            reasons.append(f"+{points} accepted role family: {role_family}")
        if rejected_count:
            points = 10 * rejected_count
            score -= points
            reasons.append(f"-{points} rejected role family: {role_family}")

    accepted_skill_hits = [
        skill for skill in matched_skills if feedback_preferences["accepted_skills"][skill]
    ]
    rejected_skill_hits = [
        skill for skill in matched_skills if feedback_preferences["rejected_skills"][skill]
    ]
    if accepted_skill_hits:
        points = 4 * sum(feedback_preferences["accepted_skills"][skill] for skill in accepted_skill_hits)
        score += points
        reasons.append(f"+{points} accepted skills: {', '.join(accepted_skill_hits[:5])}")
    if rejected_skill_hits:
        points = 4 * sum(feedback_preferences["rejected_skills"][skill] for skill in rejected_skill_hits)
        score -= points
        reasons.append(f"-{points} rejected skills: {', '.join(rejected_skill_hits[:5])}")

    if company_size:
        accepted_count = feedback_preferences["accepted_company_size"][company_size]
        rejected_count = feedback_preferences["rejected_company_size"][company_size]
        if accepted_count:
            points = 5 * accepted_count
            score += points
            reasons.append(f"+{points} accepted company size: {company_size}")
        if rejected_count:
            points = 5 * rejected_count
            score -= points
            reasons.append(f"-{points} rejected company size: {company_size}")

    if employment_type:
        rejected_count = feedback_preferences["rejected_employment_type"][employment_type]
        if rejected_count:
            points = 8 * rejected_count
            score -= points
            reasons.append(f"-{points} rejected employment type: {employment_type}")

    if bool(job.get("is_non_us_location")) or not bool(job.get("passes_location_filter", True)):
        rejected_count = feedback_preferences["rejected_hard_constraints"]["non_us_location"]
        if rejected_count:
            points = 15 * rejected_count
            score -= points
            reasons.append(f"-{points} rejected hard constraint: non-US/location mismatch")
    if bool(job.get("is_defense_military_clearance")):
        rejected_count = feedback_preferences["rejected_hard_constraints"]["security_clearance_or_defense"]
        if rejected_count:
            points = 15 * rejected_count
            score -= points
            reasons.append(f"-{points} rejected hard constraint: security/defense")
    if bool(job.get("is_contract_or_temp")):
        rejected_count = feedback_preferences["rejected_hard_constraints"]["contract_or_temp"]
        if rejected_count:
            points = 12 * rejected_count
            score -= points
            reasons.append(f"-{points} rejected hard constraint: contract/temp")

    accepted_title_hits = [
        term for term in title_terms if feedback_preferences["accepted_title_terms"][term]
    ]
    rejected_title_hits = [
        term for term in title_terms if feedback_preferences["rejected_title_terms"][term]
    ]
    if accepted_title_hits:
        points = 2 * sum(feedback_preferences["accepted_title_terms"][term] for term in accepted_title_hits)
        score += points
        reasons.append(f"+{points} accepted title terms: {', '.join(accepted_title_hits[:5])}")
    if rejected_title_hits:
        points = 2 * sum(feedback_preferences["rejected_title_terms"][term] for term in rejected_title_hits)
        score -= points
        reasons.append(f"-{points} rejected title terms: {', '.join(rejected_title_hits[:5])}")

    return score, "; ".join(reasons)


def apply_feedback_reranking(
    ranked_jobs: pd.DataFrame,
    feedback_history: list[dict[str, object]],
    protected_role_family: str = "",
) -> tuple[pd.DataFrame, dict[str, Counter]]:
    reranked_jobs = ranked_jobs.copy()
    preferences = build_feedback_preferences(feedback_history, protected_role_family)

    if not feedback_history:
        reranked_jobs["feedback_adjustment_score"] = 0
        reranked_jobs["feedback_adjustment_reason"] = ""
        reranked_jobs["adjusted_final_score"] = reranked_jobs["final_score"]
        if "feedback_label" not in reranked_jobs.columns:
            reranked_jobs["feedback_label"] = "Skip"
        return reranked_jobs, preferences

    adjustment_rows = reranked_jobs.apply(
        lambda job: score_feedback_adjustment(job, preferences),
        axis=1,
    )
    reranked_jobs["feedback_adjustment_score"] = [
        item[0] for item in adjustment_rows
    ]
    reranked_jobs["feedback_adjustment_reason"] = [
        item[1] for item in adjustment_rows
    ]
    reranked_jobs["adjusted_final_score"] = (
        reranked_jobs["final_score"] + reranked_jobs["feedback_adjustment_score"]
    ).round(2)
    label_lookup = {
        (
            str(record.get("title", "")),
            str(record.get("company", "")),
            str(record.get("location", "")),
        ): str(record.get("feedback_label", "Skip"))
        for record in feedback_history
    }
    reranked_jobs["feedback_label"] = reranked_jobs.apply(
        lambda job: label_lookup.get(
            (
                str(job.get("title", "")),
                str(job.get("company", "")),
                str(job.get("location", "")),
            ),
            "Skip",
        ),
        axis=1,
    )
    reranked_jobs = reranked_jobs.sort_values(
        by=[
            "passes_all_active_constraints",
            "adjusted_final_score",
            "role_fit_score",
            "final_score",
            "title",
        ],
        ascending=[False, False, False, False, True],
    ).reset_index(drop=True)
    reranked_jobs["rank"] = range(1, len(reranked_jobs) + 1)
    return reranked_jobs, preferences


def select_display_jobs(ranked_jobs: pd.DataFrame, advanced_preferences: dict) -> pd.DataFrame:
    display_jobs = ranked_jobs
    if (
        advanced_preferences.get("location_strictness") == "Strict"
        and "passes_location_filter" in display_jobs.columns
    ):
        display_jobs = display_jobs[display_jobs["passes_location_filter"]]

    passing_jobs = display_jobs
    if "passes_all_active_constraints" in display_jobs.columns:
        passing_jobs = display_jobs[display_jobs["passes_all_active_constraints"]]
    if not passing_jobs.empty:
        display_jobs = passing_jobs
    return display_jobs.head(10).copy()


def simulate_feedback_label(job: pd.Series, advanced_preferences: dict) -> str:
    if not bool(job.get("passes_all_active_constraints", True)):
        return "Reject"
    if avoid_contract_temp_active(advanced_preferences) and bool(job.get("is_contract_or_temp")):
        return "Reject"
    if advanced_preferences.get("company_size_preference") in {
        "Prefer large companies",
        "Large companies only",
    } and bool(job.get("is_tiny_startup")):
        return "Reject"
    if safe_numeric_or_none(job.get("role_fit_score")) is not None and float(job.get("role_fit_score")) < 0:
        return "Reject"
    if (
        bool(job.get("passes_all_active_constraints", True))
        and safe_numeric_or_none(job.get("role_fit_score")) is not None
        and float(job.get("role_fit_score")) >= 20
        and str(job.get("matched_skills", "")).strip()
    ):
        return "Accept"
    return "Skip"


def summarize_feedback_preferences(feedback_preferences: dict[str, Counter]) -> tuple[str, str]:
    net_scores: Counter = Counter()
    for accepted_key, rejected_key, accepted_weight, rejected_weight in [
        ("accepted_role_family", "rejected_role_family", 10, 10),
        ("accepted_skills", "rejected_skills", 4, 4),
        ("accepted_company_size", "rejected_company_size", 5, 5),
        ("accepted_title_terms", "rejected_title_terms", 2, 2),
    ]:
        for term, count in feedback_preferences[accepted_key].items():
            net_scores[term] += accepted_weight * count
        for term, count in feedback_preferences[rejected_key].items():
            net_scores[term] -= rejected_weight * count
    for term, count in feedback_preferences["rejected_employment_type"].items():
        net_scores[term] -= 8 * count
    for term, count in feedback_preferences["rejected_hard_constraints"].items():
        net_scores[term] -= 15 * count

    positive_summary = ", ".join(
        term for term, score in net_scores.most_common(8) if score > 0
    ) or "None yet"
    negative_summary = ", ".join(
        term
        for term, score in sorted(net_scores.items(), key=lambda item: item[1])[:8]
        if score < 0
    ) or "None yet"
    return positive_summary, negative_summary


def feedback_job_key(record: object) -> tuple[str, str, str]:
    if isinstance(record, pd.Series):
        return (
            str(record.get("title", "")),
            str(record.get("company", "")),
            str(record.get("location", "")),
        )
    if isinstance(record, dict):
        return (
            str(record.get("title", "")),
            str(record.get("company", "")),
            str(record.get("location", "")),
        )
    return ("", "", "")


def build_job_lookup(jobs: pd.DataFrame) -> dict[tuple[str, str, str], pd.Series]:
    return {feedback_job_key(row): row for _, row in jobs.iterrows()}


def average_or_none(values: list[float]) -> float | None:
    return sum(values) / len(values) if values else None


def dcg_at_k(relevance_scores: list[float], k: int = 10) -> float:
    scores = np.asarray(relevance_scores[:k], dtype=float)
    if scores.size == 0:
        return 0.0
    discounts = np.log2(np.arange(2, scores.size + 2))
    gains = np.power(2.0, scores) - 1.0
    return float(np.sum(gains / discounts))


def ndcg_at_k(relevance_scores: list[float], k: int = 10) -> float:
    actual_dcg = dcg_at_k(relevance_scores, k=k)
    ideal_scores = sorted(relevance_scores, reverse=True)
    ideal_dcg = dcg_at_k(ideal_scores, k=k)
    if ideal_dcg == 0:
        return 0.0
    return actual_dcg / ideal_dcg


def compute_feedback_ndcg(
    before_df: pd.DataFrame,
    after_df: pd.DataFrame,
    feedback_map: dict[tuple[str, str, str], str],
    k: int = 10,
) -> dict[str, float | None]:
    if not feedback_map:
        return {"before": None, "after": None, "delta": None}

    relevance_by_label = {"Accept": 2, "Skip": 1, "Reject": 0}
    before_relevance = [
        relevance_by_label.get(feedback_map.get(feedback_job_key(row), ""), 0)
        for _, row in before_df.head(k).iterrows()
    ]
    after_relevance = [
        relevance_by_label.get(feedback_map.get(feedback_job_key(row), ""), 0)
        for _, row in after_df.head(k).iterrows()
    ]
    before_ndcg = ndcg_at_k(before_relevance, k=k)
    after_ndcg = ndcg_at_k(after_relevance, k=k)
    return {
        "before": before_ndcg,
        "after": after_ndcg,
        "delta": after_ndcg - before_ndcg,
    }


def build_first_round_feedback_map(
    feedback_history: list[dict[str, object]],
) -> dict[tuple[str, str, str], str]:
    feedback_rounds = [
        int(parsed_round)
        for record in feedback_history
        if (parsed_round := safe_numeric_or_none(record.get("feedback_round"))) is not None
    ]
    if not feedback_rounds:
        return {}

    first_round = min(feedback_rounds)
    return {
        feedback_job_key(record): str(record.get("feedback_label", "Skip"))
        for record in feedback_history
        if safe_numeric_or_none(record.get("feedback_round")) == first_round
    }


def calculate_feedback_impact_metrics(
    feedback_history: list[dict[str, object]],
    baseline_jobs: pd.DataFrame,
    adjusted_jobs: pd.DataFrame,
) -> dict[str, object]:
    baseline_lookup = build_job_lookup(baseline_jobs)
    adjusted_lookup = build_job_lookup(adjusted_jobs)
    accepted_rank_before: list[float] = []
    accepted_rank_after: list[float] = []
    rejected_rank_before: list[float] = []
    rejected_rank_after: list[float] = []
    accepted_scores: list[float] = []
    rejected_scores: list[float] = []
    accepted_promoted = 0
    rejected_demoted = 0

    for record in feedback_history:
        label = record.get("feedback_label")
        if label not in {"Accept", "Reject"}:
            continue
        key = feedback_job_key(record)
        baseline_job = baseline_lookup.get(key)
        adjusted_job = adjusted_lookup.get(key)
        if baseline_job is None or adjusted_job is None:
            continue

        baseline_rank = safe_numeric_or_none(baseline_job.get("rank"))
        adjusted_rank = safe_numeric_or_none(adjusted_job.get("rank"))
        adjusted_score = safe_numeric_or_none(adjusted_job.get("adjusted_final_score"))
        if baseline_rank is None or adjusted_rank is None:
            continue

        if label == "Accept":
            accepted_rank_before.append(baseline_rank)
            accepted_rank_after.append(adjusted_rank)
            if adjusted_score is not None:
                accepted_scores.append(adjusted_score)
            if adjusted_rank < baseline_rank:
                accepted_promoted += 1
        elif label == "Reject":
            rejected_rank_before.append(baseline_rank)
            rejected_rank_after.append(adjusted_rank)
            if adjusted_score is not None:
                rejected_scores.append(adjusted_score)
            if adjusted_rank > baseline_rank:
                rejected_demoted += 1

    return {
        "accepted_avg_rank_before": average_or_none(accepted_rank_before),
        "accepted_avg_rank_after": average_or_none(accepted_rank_after),
        "rejected_avg_rank_before": average_or_none(rejected_rank_before),
        "rejected_avg_rank_after": average_or_none(rejected_rank_after),
        "accepted_promoted": accepted_promoted,
        "rejected_demoted": rejected_demoted,
        "accepted_avg_adjusted_score": average_or_none(accepted_scores),
        "rejected_avg_adjusted_score": average_or_none(rejected_scores),
    }


BATCH_ANALYTICS_SKILLS = {
    "Python": ["python"],
    "SQL": ["sql"],
    "Java": ["java"],
    "JavaScript": ["javascript"],
    "R": ["r"],
    "Excel": ["excel"],
    "Tableau": ["tableau"],
    "Power BI": ["power bi"],
    "AWS": ["aws"],
    "Azure": ["azure"],
    "GCP": ["gcp", "google cloud"],
    "Spark": ["spark"],
    "PySpark": ["pyspark"],
    "Kafka": ["kafka"],
    "Kubernetes": ["kubernetes"],
    "Docker": ["docker"],
    "TensorFlow": ["tensorflow"],
    "PyTorch": ["pytorch"],
    "Machine Learning": ["machine learning"],
    "NLP": ["nlp", "natural language processing"],
    "Data Warehouse": ["data warehouse", "data warehousing"],
    "ETL": ["etl"],
    "Airflow": ["airflow"],
    "Snowflake": ["snowflake"],
    "Databricks": ["databricks"],
    "React": ["react", "react.js", "reactjs"],
    "Node.js": ["node.js", "nodejs"],
    "C#": ["c#"],
    ".NET": [".net"],
}
SALARY_BAND_ORDER = [
    "Unknown / Not disclosed",
    "< $80K",
    "$80K-$120K",
    "$120K-$160K",
    "$160K-$200K",
    "$200K+",
]


def parse_annual_salary_for_analytics(value: object) -> float | None:
    """Estimate an annual salary while rejecting clearly non-annual values."""
    if value is None:
        return None
    try:
        if pd.isna(value):
            return None
    except (TypeError, ValueError):
        pass

    parsed_value = parse_json_like_value(value)
    if isinstance(parsed_value, dict):
        period = str(parsed_value.get("period", "")).strip().lower()
        if period and period not in {"year", "yearly", "annual", "annually"}:
            return None
        minimum = safe_numeric_or_none(parsed_value.get("minValue"))
        maximum = safe_numeric_or_none(parsed_value.get("maxValue"))
        if minimum is not None and maximum is not None:
            return (minimum + maximum) / 2
        number = safe_numeric_or_none(parsed_value.get("value"))
        if number is not None and number >= 1_000:
            return number
        value = parsed_value.get("text", "")

    text = str(value).strip()
    normalized = text.lower()
    if not text or normalized in {
        "unknown",
        "none",
        "nan",
        "null",
        "depends on experience",
        "competitive salary",
    }:
        return None
    if re.search(r"\b(?:hour|hourly|per hour|month|monthly|per month|week|weekly|day|daily)\b", normalized):
        return None

    numbers = [
        float(number.replace(",", ""))
        for number in re.findall(r"\d[\d,]*(?:\.\d+)?", text)
        if float(number.replace(",", "")) >= 1_000
    ]
    if len(numbers) >= 2:
        return (numbers[0] + numbers[1]) / 2
    if numbers:
        return numbers[0]
    return None


def salary_band(value: float | None) -> str:
    parsed_value = safe_numeric_or_none(value)
    if parsed_value is None:
        return "Unknown / Not disclosed"
    if parsed_value < 80_000:
        return "< $80K"
    if parsed_value < 120_000:
        return "$80K-$120K"
    if parsed_value < 160_000:
        return "$120K-$160K"
    if parsed_value < 200_000:
        return "$160K-$200K"
    return "$200K+"


def clean_analytics_location(value: object) -> str:
    if value is None:
        return "Unknown"
    try:
        if pd.isna(value):
            return "Unknown"
    except (TypeError, ValueError):
        pass
    location = re.sub(r"\s+", " ", str(value)).strip()
    if not location or location.lower() in {"unknown", "nan", "none", "null"}:
        return "Unknown"
    if normalize_location_text(location) in {"us", "usa", "u s", "united states"}:
        return "US unspecified"
    return location


@st.cache_data(show_spinner=False)
def compute_batch_analytics(jobs: pd.DataFrame) -> dict[str, object]:
    """Compute aggregate insights from the full cleaned job dataset."""
    total_jobs = len(jobs)
    unique_companies = (
        jobs["company"].replace("", pd.NA).nunique()
        if "company" in jobs.columns
        else 0
    )
    unique_locations = (
        jobs["location"].apply(clean_analytics_location).nunique()
        if "location" in jobs.columns
        else 0
    )

    title_text = jobs["title"].fillna("").astype(str) if "title" in jobs.columns else pd.Series("", index=jobs.index)
    description_text = jobs["description"].fillna("").astype(str) if "description" in jobs.columns else pd.Series("", index=jobs.index)
    searchable_text = (title_text + " " + description_text).str.lower()
    skill_rows: list[dict[str, object]] = []
    for skill, aliases in BATCH_ANALYTICS_SKILLS.items():
        alias_patterns = [
            rf"(?<![a-z0-9]){re.escape(alias.lower())}(?![a-z0-9])"
            for alias in aliases
        ]
        count = int(searchable_text.str.contains("|".join(alias_patterns), regex=True).sum())
        skill_rows.append({"skill": skill, "job_count": count})
    top_skills = (
        pd.DataFrame(skill_rows)
        .sort_values(["job_count", "skill"], ascending=[False, True])
        .head(15)
        .reset_index(drop=True)
    )

    salary_distribution = pd.DataFrame(columns=["band", "job_count"])
    salary_stats: dict[str, float] = {}
    known_salary_count = 0
    unknown_salary_count = total_jobs
    salary_column = "salary_text" if "salary_text" in jobs.columns else (
        "salary" if "salary" in jobs.columns else ""
    )
    if salary_column:
        parsed_salaries = jobs[salary_column].apply(parse_annual_salary_for_analytics)
        known_values = pd.to_numeric(parsed_salaries, errors="coerce").dropna()
        known_salary_count = len(known_values)
        unknown_salary_count = total_jobs - known_salary_count
        bands = parsed_salaries.apply(salary_band)
        salary_distribution = (
            bands.value_counts()
            .reindex(SALARY_BAND_ORDER, fill_value=0)
            .rename_axis("band")
            .reset_index(name="job_count")
        )
        if not known_values.empty:
            salary_stats = {
                "count": float(known_values.count()),
                "median": float(known_values.median()),
                "mean": float(known_values.mean()),
                "min": float(known_values.min()),
                "max": float(known_values.max()),
            }

    location_demand = pd.DataFrame(columns=["location", "job_count"])
    if "location" in jobs.columns:
        location_demand = (
            jobs["location"].apply(clean_analytics_location)
            .value_counts()
            .head(15)
            .rename_axis("location")
            .reset_index(name="job_count")
        )

    return {
        "total_jobs": total_jobs,
        "unique_companies": int(unique_companies),
        "unique_locations": int(unique_locations),
        "known_salary_count": known_salary_count,
        "unknown_salary_count": unknown_salary_count,
        "top_skills": top_skills,
        "salary_distribution": salary_distribution,
        "salary_stats": salary_stats,
        "location_demand": location_demand,
        "has_salary_column": bool(salary_column),
        "has_location_column": "location" in jobs.columns,
    }


def render_batch_analytics(jobs: pd.DataFrame) -> None:
    analytics = compute_batch_analytics(jobs)
    st.subheader("Batch Analytics")
    with st.expander("Full Job Dataset Insights", expanded=False):
        st.write(
            "These analytics are computed from the full cleaned job dataset, not "
            "only the current recommendations."
        )

        summary_columns = st.columns(5)
        summary_columns[0].metric("Total jobs", f"{analytics['total_jobs']:,}")
        summary_columns[1].metric("Unique companies", f"{analytics['unique_companies']:,}")
        summary_columns[2].metric("Unique locations", f"{analytics['unique_locations']:,}")
        summary_columns[3].metric("Jobs with known salary", f"{analytics['known_salary_count']:,}")
        summary_columns[4].metric("Unknown salary", f"{analytics['unknown_salary_count']:,}")

        st.subheader("Top Skills")
        top_skills = analytics["top_skills"]
        st.bar_chart(top_skills.set_index("skill")["job_count"])
        st.dataframe(top_skills, use_container_width=True, hide_index=True)
        st.download_button(
            "Download top skills CSV",
            data=top_skills.to_csv(index=False),
            file_name="jobpilot_batch_top_skills.csv",
            mime="text/csv",
            key="download_batch_top_skills",
        )

        st.subheader("Salary Distribution")
        if not analytics["has_salary_column"]:
            st.warning("Salary column is unavailable in the full job dataset.")
        else:
            salary_distribution = analytics["salary_distribution"]
            st.bar_chart(salary_distribution.set_index("band")["job_count"])
            st.dataframe(salary_distribution, use_container_width=True, hide_index=True)
            salary_stats = analytics["salary_stats"]
            if salary_stats:
                stats_columns = st.columns(5)
                stats_columns[0].metric("Parseable salaries", f"{int(salary_stats['count']):,}")
                stats_columns[1].metric("Median", f"${salary_stats['median']:,.0f}")
                stats_columns[2].metric("Mean", f"${salary_stats['mean']:,.0f}")
                stats_columns[3].metric("Min", f"${salary_stats['min']:,.0f}")
                stats_columns[4].metric("Max", f"${salary_stats['max']:,.0f}")
            st.download_button(
                "Download salary distribution CSV",
                data=salary_distribution.to_csv(index=False),
                file_name="jobpilot_batch_salary_distribution.csv",
                mime="text/csv",
                key="download_batch_salary_distribution",
            )

        st.subheader("Demand by Location")
        if not analytics["has_location_column"]:
            st.warning("Location column is unavailable in the full job dataset.")
        else:
            location_demand = analytics["location_demand"]
            st.bar_chart(location_demand.set_index("location")["job_count"])
            st.dataframe(location_demand, use_container_width=True, hide_index=True)
            st.download_button(
                "Download demand by location CSV",
                data=location_demand.to_csv(index=False),
                file_name="jobpilot_batch_location_demand.csv",
                mime="text/csv",
                key="download_batch_location_demand",
            )


def main() -> None:
    st.set_page_config(page_title="JobPilot MVP", page_icon=":briefcase:", layout="wide")

    st.title("JobPilot")
    st.caption("Smart job matcher and tailored resume builder")

    data_path = get_data_path()
    jobs = load_jobs(data_path)

    with st.sidebar:
        st.header("Search Preferences")
        uploaded_resume = st.file_uploader("Resume PDF", type=["pdf"])
        target_role = st.text_input("Target role", value="")
        skills_text = st.text_area(
            "Skills",
            value="",
            help="Separate skills with commas or new lines.",
        )
        preferred_location = st.text_input("Preferred location", value="")
        minimum_salary = st.slider(
            "Minimum salary",
            min_value=0,
            max_value=SALARY_SLIDER_MAX,
            value=0,
            step=5000,
        )
        dealbreakers_text = st.text_area(
            "Dealbreakers",
            value="",
            help="Separate dealbreakers with commas or new lines.",
        )
        with st.expander("Advanced Filters / Re-ranking Preferences"):
            role_fit_strictness = st.selectbox(
                "Role fit strictness",
                ["Flexible", "Balanced", "Strict"],
                index=1,
            )
            location_strictness = st.selectbox(
                "Location strictness",
                ["Flexible", "Balanced", "Strict"],
                index=1,
                help=(
                    "Flexible allows broader exploration, Balanced penalizes clear "
                    "mismatches, and Strict treats location preferences more like filters."
                ),
            )
            seniority_preference = st.selectbox(
                "Seniority preference",
                SENIORITY_PREFERENCE_OPTIONS,
            )
            apply_max_years_filter = st.checkbox("Apply max years filter")
            max_required_years_value = st.slider(
                "Maximum required years of experience",
                min_value=0,
                max_value=10,
                value=10,
                step=1,
            )
            max_required_years = (
                f"{max_required_years_value} years"
                if apply_max_years_filter
                else "No limit"
            )
            avoid_contract_roles = st.checkbox("Avoid contract/temp roles")
            avoid_unpaid_roles = st.checkbox("Avoid unpaid / commission-only roles")
            avoid_defense_companies = st.checkbox("Avoid defense/military companies")
            visa_sponsorship = st.selectbox(
                "Visa sponsorship preference",
                ["No preference", "Prefer sponsor-friendly", "Sponsorship required"],
            )
            company_preference = st.selectbox(
                "Company preference",
                [
                    "No preference",
                    "Prefer tech",
                    "Prefer healthcare",
                    "Prefer large companies",
                    "Avoid tiny startups",
                    "Prefer research labs",
                ],
            )
            company_size_preference = st.selectbox(
                "Company size preference",
                COMPANY_SIZE_PREFERENCE_OPTIONS,
            )
            salary_strictness = st.selectbox(
                "Salary strictness",
                ["Soft", "Medium", "Strict"],
            )
        start_matching = st.button("Start Matching", type="primary")

    seniority_flags = derive_seniority_filter_flags(seniority_preference)
    dealbreaker_terms = set(split_keywords(dealbreakers_text))
    avoid_internship_from_dealbreakers = bool(
        dealbreaker_terms
        & {"intern", "internship", "trainee", "new grad", "entry level", "entry-level"}
    )
    advanced_preferences = {
        "role_fit_strictness": role_fit_strictness,
        "location_strictness": location_strictness,
        "seniority_preference": seniority_preference,
        "avoid_senior_roles": seniority_flags["avoid_senior_roles"],
        "avoid_junior_roles": seniority_flags["avoid_junior_roles"],
        "avoid_internship_roles": avoid_internship_from_dealbreakers,
        "avoid_manager_roles": seniority_flags["avoid_manager_roles"],
        "apply_max_years_filter": apply_max_years_filter,
        "max_required_years": max_required_years,
        "max_required_years_value": max_required_years_value if apply_max_years_filter else None,
        "minimum_salary": minimum_salary,
        "avoid_contract_roles": avoid_contract_roles,
        "avoid_contract_temp": avoid_contract_roles,
        "avoid_unpaid_roles": avoid_unpaid_roles,
        "avoid_unpaid_commission": avoid_unpaid_roles,
        "avoid_defense_companies": avoid_defense_companies,
        "visa_sponsorship": visa_sponsorship,
        "company_preference": company_preference,
        "company_size_preference": company_size_preference,
        "salary_strictness": salary_strictness,
    }

    resume_text = ""
    candidate_name = "Unknown"
    extracted_resume_skills: list[str] = []
    if uploaded_resume is not None:
        try:
            resume_text = extract_text_from_pdf(uploaded_resume)
            candidate_name = clean_candidate_name(extract_candidate_name(resume_text))
            extracted_resume_skills = extract_skills(resume_text)
        except Exception as error:
            st.warning(f"Could not extract text from this PDF: {error}")

    combined_skills_text = combine_skill_inputs(skills_text, extracted_resume_skills)

    if uploaded_resume is not None and resume_text:
        st.caption("Resume uploaded and parsed successfully.")

    if combined_skills_text:
        display_profile_signals = format_profile_signals_for_display(combined_skills_text)
        if display_profile_signals:
            st.subheader("Profile Signals Used for Matching")
            st.write(display_profile_signals)

    if start_matching:
        retrieval_used = False
        retrieval_note = "Used full rule-based ranking."
        candidate_jobs = jobs
        try:
            candidate_jobs, retrieval_used, retrieval_note = get_embedding_candidates(
                jobs=jobs,
                target_role=target_role,
                combined_skills_text=combined_skills_text,
                resume_text=resume_text,
            )
        except Exception as error:
            st.warning(f"Embedding retrieval unavailable, using rule-based fallback: {error}")

        ranked_jobs = rank_jobs(
            jobs=candidate_jobs,
            target_role=target_role,
            skills_text=combined_skills_text,
            preferred_location=preferred_location,
            minimum_salary=minimum_salary,
            dealbreakers_text=dealbreakers_text,
            advanced_preferences=advanced_preferences,
        )
        if retrieval_used:
            ranked_jobs["explanation"] = (
                "Retrieved by semantic similarity, then ranked by title, skill, "
                "location, salary, and dealbreaker rules. "
                + ranked_jobs["explanation"]
            )
        if "embedding_similarity" not in ranked_jobs.columns:
            ranked_jobs["embedding_similarity"] = 0.0
        st.session_state["baseline_top_10_jobs"] = select_display_jobs(
            ranked_jobs,
            advanced_preferences,
        )
        st.session_state["baseline_ranked_jobs"] = ranked_jobs
        st.session_state["top_10_jobs"] = st.session_state["baseline_top_10_jobs"]
        st.session_state["full_ranked_jobs"] = ranked_jobs
        st.session_state["feedback_adjusted_jobs"] = ranked_jobs
        st.session_state["feedback_preferences"] = build_feedback_preferences(
            st.session_state.get("feedback_history", []),
            detect_role_family(target_role),
        )
        st.session_state["matched_target_role_family"] = detect_role_family(target_role)
        st.session_state["matched_target_role"] = target_role
        st.session_state["matched_candidate_name"] = clean_candidate_name(candidate_name)
        st.session_state["matched_resume_text"] = resume_text
        st.session_state["combined_skills_text"] = combined_skills_text
        st.session_state["matched_data_path"] = str(data_path)
        st.session_state["matched_preferred_location"] = preferred_location
        st.session_state["retrieval_used"] = retrieval_used
        st.session_state["retrieval_note"] = retrieval_note
        st.session_state["advanced_preferences"] = advanced_preferences
        st.session_state["has_matched"] = True

    if not st.session_state.get("has_matched"):
        st.info("Enter your profile preferences and click Start Matching to see ranked jobs.")
        render_batch_analytics(jobs)
        return

    top_10 = st.session_state["top_10_jobs"]
    full_ranked_jobs = st.session_state.get("full_ranked_jobs", top_10)
    matched_data_path = st.session_state.get("matched_data_path", str(data_path))
    matched_skills_text = st.session_state.get("combined_skills_text", combined_skills_text)
    matched_target_role = st.session_state.get("matched_target_role", target_role)
    matched_candidate_name = clean_candidate_name(
        st.session_state.get("matched_candidate_name", candidate_name)
    )
    matched_resume_text = st.session_state.get("matched_resume_text", resume_text)
    retrieval_used = st.session_state.get("retrieval_used", False)
    retrieval_note = st.session_state.get("retrieval_note", "Used full rule-based ranking.")
    matched_preferred_location = st.session_state.get(
        "matched_preferred_location",
        preferred_location,
    )
    active_advanced_preferences = st.session_state.get(
        "advanced_preferences",
        advanced_preferences,
    )
    matched_target_role_family = st.session_state.get(
        "matched_target_role_family",
        detect_role_family(target_role),
    )
    st.session_state.setdefault("feedback_history", [])
    st.session_state.setdefault("feedback_round", 0)
    if (
        active_advanced_preferences.get("location_strictness") == "Strict"
        and "passes_location_filter" in top_10.columns
    ):
        top_10 = top_10[top_10["passes_location_filter"]].copy()
        st.session_state["top_10_jobs"] = top_10

    st.subheader("Top Ranked Jobs")
    with st.expander("How recommendations and Match Score work", expanded=False):
        st.write(
            "JobPilot retrieves semantically similar jobs using ANN over job embeddings, "
            "then re-ranks them using role fit, matched skills, location/salary "
            "preferences, constraints, and adaptive feedback."
        )
        st.write(
            "Match Score is the current ranking score used to order recommendations. "
            "It combines semantic embedding similarity, rule-based fit, matched skills, "
            "location/salary preferences, seniority/company constraints, and adaptive "
            "feedback adjustments."
        )
        st.write(
            "Higher Match Score means a stronger predicted fit for the candidate "
            "profile. Before feedback, Match Score is usually the base final score. "
            "After feedback, Match Score includes feedback-based boosts and penalties."
        )
    if len(top_10) < 10:
        st.info(
            f"Only {len(top_10)} jobs passed all active constraints. Strict filters "
            "may return fewer than 10 results rather than filling the list with "
            "lower-quality matches."
        )

    display_jobs = top_10.copy()
    if "adjusted_final_score" in display_jobs.columns:
        display_jobs["Match Score"] = pd.to_numeric(
            display_jobs["adjusted_final_score"],
            errors="coerce",
        )
    else:
        display_jobs["Match Score"] = pd.NA
    if "final_score" in display_jobs.columns:
        display_jobs["Match Score"] = display_jobs["Match Score"].fillna(
            pd.to_numeric(display_jobs["final_score"], errors="coerce")
        )
    display_jobs["Match Score"] = display_jobs["Match Score"].round(2)

    display_columns = [
        "rank",
        "title",
        "company",
        "location",
        "salary_text",
        "required_years_min",
        "employment_type_inferred",
        "company_size_inferred",
        "matched_skills",
        "Match Score",
    ]
    display_column_names = {
        "rank": "Rank",
        "title": "Job Title",
        "company": "Company",
        "location": "Location",
        "salary_text": "Salary",
        "required_years_min": "Required Years",
        "employment_type_inferred": "Employment Type",
        "company_size_inferred": "Company Size",
        "matched_skills": "Matched Skills",
        "Match Score": "Match Score",
    }
    display_table = display_jobs[
        existing_columns(display_jobs, display_columns)
    ].copy()
    for column in display_table.columns:
        if column == "Match Score":
            continue
        display_table[column] = display_table[column].apply(
            lambda value: "Unknown"
            if value is None
            or pd.isna(value)
            or str(value).strip().lower() in {"", "nan", "none", "unknown"}
            else value
        )
    display_table = display_table.rename(columns=display_column_names)
    st.dataframe(
        display_table,
        use_container_width=True,
        hide_index=True,
    )

    st.subheader("Adaptive Feedback")
    st.caption(
        "Review the recommendations below. Your Accept / Reject / Skip choices help "
        "JobPilot learn your preferences and re-rank jobs to better match your profile."
    )

    feedback_options = ["Skip", "Accept", "Reject"]
    for _, job in top_10.iterrows():
        feedback_key = (
            f"feedback_{job.get('rank', '')}_"
            f"{str(job.get('title', '')).replace(' ', '_')[:40]}_"
            f"{str(job.get('company', '')).replace(' ', '_')[:30]}"
        )
        current_label = str(job.get("feedback_label", "Skip"))
        current_index = feedback_options.index(current_label) if current_label in feedback_options else 0
        st.selectbox(
            f"{job['rank']}. {job['title']} at {job['company']}",
            feedback_options,
            index=current_index,
            key=feedback_key,
        )

    feedback_col_1, feedback_col_2 = st.columns(2)
    with feedback_col_1:
        apply_feedback = st.button("Apply Feedback and Re-rank")
    with feedback_col_2:
        reset_feedback = st.button("Reset Feedback")
    simulate_feedback = False

    if reset_feedback:
        st.session_state["feedback_history"] = []
        st.session_state["feedback_round"] = 0
        base_jobs = st.session_state.get("baseline_ranked_jobs", full_ranked_jobs).copy()
        base_jobs["feedback_label"] = "Skip"
        base_jobs["feedback_adjustment_score"] = 0
        base_jobs["feedback_adjustment_reason"] = ""
        base_jobs["adjusted_final_score"] = base_jobs["final_score"]
        st.session_state["full_ranked_jobs"] = base_jobs
        st.session_state["feedback_adjusted_jobs"] = base_jobs
        st.session_state["top_10_jobs"] = select_display_jobs(
            base_jobs,
            active_advanced_preferences,
        )
        st.session_state["feedback_preferences"] = build_feedback_preferences(
            [],
            matched_target_role_family,
        )
        st.rerun()

    if simulate_feedback:
        next_round = int(st.session_state.get("feedback_round", 0)) + 1
        simulated_records = [
            build_feedback_record(
                job,
                next_round,
                simulate_feedback_label(job, active_advanced_preferences),
            )
            for _, job in top_10.iterrows()
        ]
        st.session_state["feedback_round"] = next_round
        st.session_state["feedback_history"].extend(simulated_records)
        adjusted_jobs, learned_preferences = apply_feedback_reranking(
            st.session_state.get("baseline_ranked_jobs", full_ranked_jobs),
            st.session_state["feedback_history"],
            matched_target_role_family,
        )
        st.session_state["feedback_adjusted_jobs"] = adjusted_jobs
        st.session_state["full_ranked_jobs"] = adjusted_jobs
        st.session_state["feedback_preferences"] = learned_preferences
        st.session_state["top_10_jobs"] = select_display_jobs(
            adjusted_jobs,
            active_advanced_preferences,
        )
        st.rerun()

    if apply_feedback:
        next_round = int(st.session_state.get("feedback_round", 0)) + 1
        feedback_records: list[dict[str, object]] = []
        for _, job in top_10.iterrows():
            feedback_key = (
                f"feedback_{job.get('rank', '')}_"
                f"{str(job.get('title', '')).replace(' ', '_')[:40]}_"
                f"{str(job.get('company', '')).replace(' ', '_')[:30]}"
            )
            feedback_records.append(
                build_feedback_record(
                    job,
                    next_round,
                    st.session_state.get(feedback_key, "Skip"),
                )
            )
        st.session_state["feedback_round"] = next_round
        st.session_state["feedback_history"].extend(feedback_records)
        adjusted_jobs, learned_preferences = apply_feedback_reranking(
            st.session_state.get("baseline_ranked_jobs", full_ranked_jobs),
            st.session_state["feedback_history"],
            matched_target_role_family,
        )
        st.session_state["feedback_adjusted_jobs"] = adjusted_jobs
        st.session_state["full_ranked_jobs"] = adjusted_jobs
        st.session_state["feedback_preferences"] = learned_preferences
        st.session_state["top_10_jobs"] = select_display_jobs(
            adjusted_jobs,
            active_advanced_preferences,
        )
        st.rerun()

    feedback_history = st.session_state.get("feedback_history", [])
    feedback_preferences = st.session_state.get(
        "feedback_preferences",
        build_feedback_preferences(feedback_history, matched_target_role_family),
    )
    positive_summary, negative_summary = summarize_feedback_preferences(
        feedback_preferences
    )
    accepted_count = sum(1 for record in feedback_history if record.get("feedback_label") == "Accept")
    rejected_count = sum(1 for record in feedback_history if record.get("feedback_label") == "Reject")
    skipped_count = sum(1 for record in feedback_history if record.get("feedback_label") == "Skip")
    baseline_top_10 = st.session_state.get("baseline_top_10_jobs", top_10)
    baseline_ranked_jobs = st.session_state.get("baseline_ranked_jobs", baseline_top_10)
    adjusted_ranked_jobs = st.session_state.get("feedback_adjusted_jobs", full_ranked_jobs)
    baseline_avg = pd.to_numeric(
        baseline_top_10.get("final_score", pd.Series(dtype=float)),
        errors="coerce",
    ).mean()
    current_avg = pd.to_numeric(
        top_10.get("adjusted_final_score", top_10.get("final_score", pd.Series(dtype=float))),
        errors="coerce",
    ).mean()
    impact_metrics = calculate_feedback_impact_metrics(
        feedback_history,
        baseline_ranked_jobs,
        adjusted_ranked_jobs,
    )
    first_round_feedback_map = build_first_round_feedback_map(feedback_history)
    feedback_ndcg = compute_feedback_ndcg(
        baseline_top_10,
        top_10,
        first_round_feedback_map,
        k=10,
    )

    def format_metric(value: object) -> str:
        return f"{value:.1f}" if isinstance(value, float) else "N/A"

    def format_ndcg(value: object) -> str:
        return f"{value:.3f}" if isinstance(value, float) else "N/A"

    st.write(
        f"**Feedback round:** {st.session_state.get('feedback_round', 0)} | "
        f"Accepted: {accepted_count} | Rejected: {rejected_count} | Skipped: {skipped_count}"
    )
    metric_col_1, metric_col_2, metric_col_3, metric_col_4 = st.columns(4)
    metric_col_1.metric(
        "Accepted avg rank",
        f"{format_metric(impact_metrics['accepted_avg_rank_before'])} -> {format_metric(impact_metrics['accepted_avg_rank_after'])}",
    )
    metric_col_2.metric(
        "Rejected avg rank",
        f"{format_metric(impact_metrics['rejected_avg_rank_before'])} -> {format_metric(impact_metrics['rejected_avg_rank_after'])}",
    )
    metric_col_3.metric("Accepted jobs promoted", str(impact_metrics["accepted_promoted"]))
    metric_col_4.metric("Rejected jobs demoted", str(impact_metrics["rejected_demoted"]))
    ndcg_col_1, ndcg_col_2, ndcg_col_3 = st.columns(3)
    ndcg_col_1.metric("NDCG@10 before", format_ndcg(feedback_ndcg["before"]))
    ndcg_col_2.metric("NDCG@10 after", format_ndcg(feedback_ndcg["after"]))
    ndcg_col_3.metric("Δ NDCG@10", format_ndcg(feedback_ndcg["delta"]))
    st.caption("NDCG@10 uses Accept=2, Skip=1, Reject=0 as proxy relevance labels.")

    user_download_jobs = prepare_user_download_jobs(top_10)
    user_download_columns = [
        "rank",
        "title",
        "company",
        "location",
        "salary_text",
        "description",
        "apply_url",
        "final_score",
        "matched_skills",
        "match_explanation",
    ]

    st.subheader("Job Details")
    for _, job in top_10.iterrows():
        match_score = safe_numeric_or_none(job.get("adjusted_final_score"))
        if match_score is None:
            match_score = safe_numeric_or_none(job.get("final_score")) or 0
        with st.expander(
            f"{match_score:.2f} Match Score | {job['title']} at {job['company']}"
        ):
            st.markdown("#### Job Summary")
            st.write(f"**Title:** {job.get('title', 'Unknown')}")
            st.write(f"**Company:** {job.get('company', 'Unknown')}")
            st.write(f"**Location:** {job.get('location', 'Unknown')}")
            st.write(f"**Salary:** {job.get('salary_text', 'Unknown')}")
            matched_job_skills = str(job.get("matched_skills", "")).strip()
            if matched_job_skills:
                st.write(f"**Matched Skills:** {matched_job_skills}")
            company_size = str(job.get("company_size_inferred", "")).strip()
            if company_size and company_size.lower() != "unknown":
                st.write(f"**Company Size:** {company_size}")
            employment_type = str(job.get("employment_type_inferred", "")).strip()
            if employment_type and employment_type.lower() != "unknown":
                st.write(f"**Employment Type:** {employment_type}")
            required_years = job.get("required_years_min", "Unknown")
            if str(required_years).strip().lower() not in {"", "unknown", "nan", "none"}:
                st.write(f"**Required Years:** {required_years}")
            apply_link = str(job.get("apply_link", "")).strip()
            if apply_link:
                st.write(f"**Apply link:** {apply_link}")

            st.markdown("#### Why this ranked here")
            explanation_parts = [
                part.strip()
                for part in re.split(r";\s*", str(job.get("explanation", "")))
                if part.strip()
            ]
            if explanation_parts:
                for explanation_part in explanation_parts:
                    st.markdown(f"- {explanation_part}")
            else:
                st.write("No ranking explanation is available.")

            st.markdown("#### Actions")
            if apply_link:
                st.link_button("Apply", apply_link)

            resume_key = stable_widget_key(
                "tailored_resume",
                job.get("rank", ""),
                job.get("title", ""),
                job.get("company", ""),
            )
            if st.button(
                "Generate Resume for this job",
                key=f"generate_{resume_key}",
            ):
                st.session_state[resume_key] = generate_tailored_resume(
                    candidate_name=matched_candidate_name,
                    resume_text=matched_resume_text,
                    selected_job=job,
                    matched_skills=str(job.get("matched_skills", "")),
                    target_role=matched_target_role,
                    combined_skills_text=matched_skills_text,
                )

            if resume_key in st.session_state:
                tailored_resume = st.session_state[resume_key]
                draft_key = f"draft_{resume_key}"
                edited_resume = st.text_area(
                    "Tailored resume draft",
                    value=tailored_resume,
                    height=500,
                    key=draft_key,
                )
                st.session_state[resume_key] = edited_resume
                try:
                    resume_docx = build_resume_docx(edited_resume)
                    cleaned_filename_name = clean_candidate_name(matched_candidate_name)
                    filename_name = (
                        cleaned_filename_name
                        if cleaned_filename_name != "[Add your name]"
                        else "candidate"
                    )
                    download_filename = stable_widget_key(
                        "tailored_resume",
                        filename_name,
                        job.get("title", ""),
                    ).lower()
                    download_filename = re.sub(
                        r"_+",
                        "_",
                        re.sub(r"[^a-z0-9]+", "_", download_filename),
                    ).strip("_")
                    st.download_button(
                        "Download tailored resume as DOCX",
                        data=resume_docx,
                        file_name=f"{download_filename}.docx",
                        mime=(
                            "application/vnd.openxmlformats-officedocument."
                            "wordprocessingml.document"
                        ),
                        key=f"download_docx_{resume_key}",
                    )
                except ImportError:
                    st.warning(
                        "DOCX export requires the `python-docx` package. "
                        "Install the project requirements to enable this download."
                    )

            with st.expander("Job Description", expanded=False):
                st.write(str(job.get("description", "")))

    st.subheader("Downloads")
    st.download_button(
        label="Download top 10 jobs as CSV",
        data=user_download_jobs[
            existing_columns(user_download_jobs, user_download_columns)
        ].to_csv(index=False),
        file_name="jobpilot_top_10_jobs.csv",
        mime="text/csv",
    )

    render_batch_analytics(jobs)

    st.subheader("Developer / Debug Tools")
    with st.expander("System Details / Debug", expanded=False):
        st.write(f"**Loaded job data:** `{matched_data_path}`")
        st.write(f"**Retrieval note:** {retrieval_note}")
        st.write(
            f"**Retrieval method:** "
            f"{'ANN vector search' if retrieval_used else 'Brute-force/rule-based fallback'}"
        )
        if matched_skills_text:
            st.write(f"**Combined skills used internally:** {matched_skills_text}")
        st.write(f"**Learned positive signals:** {positive_summary}")
        st.write(f"**Learned negative signals:** {negative_summary}")

        debug_metric_col_1, debug_metric_col_2 = st.columns(2)
        debug_metric_col_1.metric(
            "Round 0 avg final_score",
            f"{baseline_avg:.2f}" if pd.notna(baseline_avg) else "N/A",
        )
        debug_metric_col_2.metric(
            "Current avg adjusted score",
            f"{current_avg:.2f}" if pd.notna(current_avg) else "N/A",
        )
        debug_metric_col_3, debug_metric_col_4 = st.columns(2)
        debug_metric_col_3.metric(
            "Accepted avg adjusted score",
            format_metric(impact_metrics["accepted_avg_adjusted_score"]),
        )
        debug_metric_col_4.metric(
            "Rejected avg adjusted score",
            format_metric(impact_metrics["rejected_avg_adjusted_score"]),
        )
        st.write(
            "**Location display:** Job posting work location when available, not "
            "company headquarters."
        )
        st.download_button(
            label="Download full debug results CSV",
            data=full_ranked_jobs.to_csv(index=False),
            file_name="jobpilot_full_debug_results.csv",
            mime="text/csv",
            key="download_full_debug_results",
        )

    with st.expander("Debug: constraint checks", expanded=False):
        st.write(f"**Preferred location:** {matched_preferred_location or 'None entered'}")
        st.write(
            f"**Location strictness:** "
            f"{active_advanced_preferences.get('location_strictness', 'Balanced')}"
        )
        debug_columns = [
            "rank",
            "title",
            "company",
            "location",
            "preferred_location",
            "location_strictness",
            "seniority_preference",
            "salary_minimum",
            "salary_strictness",
            "max_required_years_setting",
            "company_size_preference",
            "company_size_inferred",
            "company_size_confidence",
            "company_size_match_text",
            "is_tiny_startup",
            "is_small_company",
            "is_medium_company",
            "is_large_company",
            "is_enterprise_company",
            "passes_company_size_filter",
            "location_source",
            "location_confidence",
            "is_us_location",
            "is_remote",
            "is_us_remote_location",
            "has_explicit_non_us_location_signal",
            "is_hybrid",
            "is_bay_area",
            "is_non_us_location",
            "passes_location_filter",
            "location_filter_reason",
            "required_years_min",
            "years_requirement_text",
            "passes_years_filter",
            "employment_type_inferred",
            "is_contract_or_temp",
            "is_unpaid_or_commission",
            "passes_contract_filter",
            "passes_unpaid_filter",
            "passes_salary_filter",
            "active_constraint_filters",
            "is_senior_staff_principal",
            "is_junior_entry_level",
            "is_manager_director",
            "passes_seniority_filter",
            "passes_internship_filter",
            "is_defense_military_clearance",
            "defense_match_text",
            "passes_defense_filter",
            "has_visa_sponsorship_signal",
            "has_no_sponsorship_signal",
            "role_fit_score",
            "passes_role_fit_filter",
            "embedding_similarity",
            "feedback_label",
            "feedback_adjustment_score",
            "adjusted_final_score",
            "passes_all_active_constraints",
        ]
        st.dataframe(
            top_10[existing_columns(top_10, debug_columns)],
            use_container_width=True,
            hide_index=True,
        )


if __name__ == "__main__":
    main()
